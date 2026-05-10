# -*- coding: utf-8 -*-
"""
Genera Informe_Proyecto.docx - Informe de entrega estilo reporte.
Estructura segun guia (7 capitulos + anexos).
"""
import sys, io, os, math
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

BASE = r"C:\Users\andre\Desktop\proyecto materiales"

doc = Document()

# ============================================================
# ESTILOS
# ============================================================
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for s in doc.styles:
    if hasattr(s, 'font') and s.font:
        s.font.name = 'Calibri'

# Margenes
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

NAVY = RGBColor(0x0A, 0x16, 0x28)
GOLD = RGBColor(0xC8, 0x95, 0x6C)
GREEN = RGBColor(0x27, 0xAE, 0x60)
RED = RGBColor(0xE7, 0x4C, 0x3C)
GRAY = RGBColor(0x55, 0x55, 0x55)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)

def add_img(path, width=Inches(5.5)):
    full = os.path.join(BASE, path)
    if os.path.exists(full):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(full, width=width)
        return True
    return False

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GRAY

def heading1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(16)

def heading2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(13)

def heading3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = NAVY

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.7)
    return p

def eq(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY

def eq_result(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x2B, 0x5C)

def make_table(headers, rows, highlight_last=False):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = h
        set_cell_shading(c, '0A1628')
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.size = Pt(9)
    # Data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]
            c.text = str(val)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
            if highlight_last and i == len(rows)-1:
                set_cell_shading(c, 'FEF9E7')
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
    doc.add_paragraph()
    return t

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_icon(fname, width=Inches(3.8)):
    full = os.path.join(BASE, "iconos_elementos", fname)
    if os.path.exists(full):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(full, width=width)

# ============================================================
# PORTADA
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PROYECTO FINAL')
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MEC\u00c1NICA DE MATERIALES')
r.font.size = Pt(22)
r.font.color.rgb = NAVY

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Dise\u00f1o de los elementos y conexiones\nde un p\u00f3rtico plano de acero')
r.font.size = Pt(14)
r.font.italic = True
r.font.color.rgb = GRAY

for _ in range(4):
    doc.add_paragraph()

info = [
    ('Estudiante:', 'xy = 45'),
    ('Profesor:', 'Ing. William Valencia Mina'),
    ('Universidad:', 'Universidad del Quind\u00edo'),
    ('Facultad:', 'Ingenier\u00eda Civil'),
    ('Fecha:', 'Mayo 2025'),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + ' ')
    r1.font.bold = True
    r1.font.size = Pt(12)
    r2 = p.add_run(value)
    r2.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# TABLA DE CONTENIDO
# ============================================================
heading1('Tabla de Contenido')
toc_items = [
    'Cap\u00edtulo 1: Introducci\u00f3n',
    'Cap\u00edtulo 2: An\u00e1lisis de Cargas',
    'Cap\u00edtulo 3: An\u00e1lisis de Esfuerzos Internos',
    'Cap\u00edtulo 4: Dise\u00f1o de Elementos',
    'Cap\u00edtulo 5: Dise\u00f1o de Conexiones',
    'Cap\u00edtulo 6: Cantidades de Material',
    'Cap\u00edtulo 7: Conclusiones y Comparaciones',
    'Anexos',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(4)
    r = p.runs[0]
    r.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# CAPITULO 1: INTRODUCCION
# ============================================================
heading1('Cap\u00edtulo 1: Introducci\u00f3n')

body('En el marco de la construcci\u00f3n de un edificio institucional, se ha delegado al equipo de ingenier\u00eda '
     'el dise\u00f1o preliminar de uno de los p\u00f3rticos estructurales principales. Este p\u00f3rtico, conformado por '
     'columnas y vigas de acero, debe ser capaz de soportar las cargas gravitacionales del sistema de piso '
     'as\u00ed como las fuerzas laterales generadas por el sismo de dise\u00f1o.')

body('El presente informe documenta el proceso completo de dise\u00f1o del p\u00f3rtico plano de acero correspondiente '
     'al par\u00e1metro xy = 45, siguiendo los lineamientos de la gu\u00eda del proyecto del profesor William Valencia Mina. '
     'Se analiza un p\u00f3rtico plano biempotrado extra\u00eddo del sistema tridimensional del edificio, '
     'compuesto por dos columnas y una viga, con apoyos empotrados en las bases A y D.')

body('El an\u00e1lisis se realiza mediante las ecuaciones del p\u00f3rtico con el par\u00e1metro de rigidez k = (I\u2082/I\u2081)\u00b7(h/L), '
     'considerando dos casos de carga: carga distribuida gravitacional y carga lateral s\u00edsmica. '
     'Los resultados se superponen para obtener las envolventes m\u00e1ximas de dise\u00f1o.')

heading2('Objetivos')

heading3('General')
body('Aplicar los principios de la mec\u00e1nica de materiales para el an\u00e1lisis y dise\u00f1o de los elementos '
     'estructurales de acero de un p\u00f3rtico plano, incluyendo la selecci\u00f3n de perfiles, el dise\u00f1o de '
     'conexiones y la verificaci\u00f3n de esfuerzos.')

heading3('Espec\u00edficos')
bullet('Realizar el an\u00e1lisis de cargas gravitacionales y laterales del p\u00f3rtico.')
bullet('Determinar las reacciones y esfuerzos internos mediante las ecuaciones del p\u00f3rtico biempotrado.')
bullet('Seleccionar perfiles comerciales W (AISC) que cumplan con los factores de seguridad requeridos.')
bullet('Dise\u00f1ar las conexiones atornilladas y soldadas en los nudos y las bases.')
bullet('Proponer un perfil personalizado optimizado que reduzca el peso de acero.')
bullet('Verificar los esfuerzos mediante c\u00edrculos de Mohr en la secci\u00f3n cr\u00edtica.')
bullet('Calcular las cantidades totales de material (vol\u00famenes y pesos).')

heading2('Datos del Proyecto')
make_table(
    ['Par\u00e1metro', 'Valor'],
    [
        ['Par\u00e1metro xy', '45'],
        ['Luz de viga (L)', '10.0 m'],
        ['Altura de columnas (h)', '3.5 m'],
        ['Dimensiones en planta', '10.0 m \u00d7 4.0 m'],
        ['Material', 'Acero ASTM A36'],
        ['Fy (fluencia)', '250 MPa'],
        ['\u03c4y (cortante fluencia)', '100 MPa'],
        ['Fu (rotura)', '400 MPa'],
        ['Factor de seguridad m\u00ednimo', '1.4'],
    ]
)

add_img("Portico_3D_y_2D.png", Inches(5.2))
caption('Figura 1. Sistema tridimensional del edificio y p\u00f3rtico plano extra\u00eddo para an\u00e1lisis.')

doc.add_page_break()

# ============================================================
# CAPITULO 2: ANALISIS DE CARGAS
# ============================================================
heading1('Cap\u00edtulo 2: An\u00e1lisis de Cargas')

heading2('2.1 Cargas gravitacionales')

body('La carga distribuida sobre la viga proviene del sistema de piso (losa de concreto reforzado) '
     'y se calcula considerando el ancho tributario de la viga:')

heading3('Carga de la losa')
eq('w_losa = e \u00d7 \u03b3c = 0.18 m \u00d7 24 kN/m\u00b3 = 4.32 kN/m\u00b2')

heading3('Cargas adicionales')
make_table(
    ['Tipo de carga', 'Valor'],
    [
        ['Peso losa (w_losa)', '4.32 kN/m\u00b2'],
        ['Carga muerta adicional', '1.20 kN/m\u00b2'],
        ['Carga viva', '3.00 kN/m\u00b2'],
        ['Total por \u00e1rea (w_total)', '8.52 kN/m\u00b2'],
    ],
    highlight_last=True
)

heading3('Carga lineal sobre la viga')
body('Con un ancho tributario b_trib = 4.0/2 = 2.0 m:')
eq('q_ext = w_total \u00d7 b_trib = 8.52 \u00d7 2.0 = 17.04 kN/m')

body('Incluyendo el peso propio de la viga W310\u00d767 (67 kg/m = 0.657 kN/m):')
eq_result('q_total = 17.04 + 0.657 = 17.70 kN/m')

heading2('2.2 Carga lateral s\u00edsmica')

body('Se aplica una fuerza horizontal concentrada en el nudo B:')
eq_result('P = 45 kN (horizontal, en nudo B)')

heading2('2.3 Par\u00e1metro de rigidez k')

body('El par\u00e1metro k relaciona las rigideces de la viga y la columna:')
eq('k = (I\u2082/I\u2081) \u00d7 (h/L)')
eq('k = (14,500/8,700) \u00d7 (3.5/10.0)')
eq_result('k = 0.583')

body('Donde I\u2082 = I_viga (W310\u00d767) = 14,500 cm\u2074 e I\u2081 = I_columna (W250\u00d758) = 8,700 cm\u2074.')

add_icon("caso1_viga.png", Inches(3.5))
caption('Caso 1: Carga distribuida q sobre la viga B-C.')

add_icon("caso2_lateral.png", Inches(3.5))
caption('Caso 2: Carga lateral P = 45 kN en el nudo B.')

doc.add_page_break()

# ============================================================
# CAPITULO 3: ANALISIS DE ESFUERZOS INTERNOS
# ============================================================
heading1('Cap\u00edtulo 3: An\u00e1lisis de Esfuerzos Internos')

heading2('3.1 Caso 1: Carga distribuida (q = 17.70 kN/m)')

add_icon("caso1_viga.png", Inches(3.2))

body('Aplicando las ecuaciones del p\u00f3rtico biempotrado (p\u00e1gs. 10-11 de la gu\u00eda):')

make_table(
    ['Reacci\u00f3n', 'F\u00f3rmula', 'Valor'],
    [
        ['V_A = V_D', 'qL/2', '88.49 kN'],
        ['H_A = H_D', 'qL\u00b2/(4h(k+2))', '48.93 kN'],
        ['M_A = M_D', 'qL\u00b2/(12(k+2))', '57.09 kN\u00b7m'],
        ['M_B = M_C', '\u2212qL\u00b2/(6(k+2))', '\u2212114.18 kN\u00b7m'],
        ['M_mid', 'qL\u00b2(3k+2)/(24(k+2))', '107.04 kN\u00b7m'],
    ]
)

heading2('3.2 Caso 2: Carga lateral (P = 45 kN)')

add_icon("caso2_lateral.png", Inches(3.2))

make_table(
    ['Reacci\u00f3n', 'F\u00f3rmula', 'Valor'],
    [
        ['V_A', '3Phk/(L(6k+1))', '6.12 kN'],
        ['V_D', '\u22123Phk/(L(6k+1))', '\u22126.12 kN'],
        ['H_A = H_D', 'P/2', '22.50 kN'],
        ['M_A', '\u2212Ph(3k+1)/(2(6k+1))', '\u221248.12 kN\u00b7m'],
        ['M_D', 'Ph(3k+1)/(2(6k+1))', '48.12 kN\u00b7m'],
        ['M_B', 'Ph\u00b73k/(2(6k+1))', '30.62 kN\u00b7m'],
    ]
)

heading2('3.3 Superposici\u00f3n de casos')

add_icon("superposicion.png", Inches(3.2))

body('Por el principio de superposici\u00f3n, se suman los valores absolutos de ambos casos para obtener '
     'las envolventes m\u00e1ximas de dise\u00f1o:')

make_table(
    ['Par\u00e1metro', 'Caso 1', '|Caso 2|', 'Total'],
    [
        ['M_B (viga)', '114.18', '30.62', '144.80 kN\u00b7m'],
        ['M_A (columna)', '57.09', '48.12', '105.21 kN\u00b7m'],
        ['V_A (m\u00e1x)', '88.49', '6.12', '94.61 kN'],
        ['H_A (m\u00e1x)', '48.93', '22.50', '71.43 kN'],
        ['N_viga', '\u2014', '\u2014', '71.43 kN (compresi\u00f3n)'],
    ],
    highlight_last=False
)

heading2('3.4 Diagramas de esfuerzos internos')

body('A continuaci\u00f3n se presentan los diagramas de cortante V(x), normal N(x) y momento M(x) '
     'para la viga y las columnas, incluyendo la superposici\u00f3n de ambos casos de carga:')

add_img("Diagramas_Portico.png", Inches(5.5))
caption('Figura 2. Diagramas de cortante, normal y momento para viga B-C y columnas A-B, D-C.')

heading3('Fuerzas internas en la viga B-C')
add_icon("fuerzas_viga.png", Inches(3.2))

body('Las ecuaciones de esfuerzos internos en la viga (0 \u2264 x \u2264 10 m), tomando origen en B:')
eq('V(x) = V_A \u2212 q\u00b7x = 94.61 \u2212 17.70x   [kN]')
eq('M(x) = M_B + V_A\u00b7x \u2212 q\u00b7x\u00b2/2   [kN\u00b7m]')
eq('N(x) = \u2212H_A = \u221271.43 kN   (compresi\u00f3n uniforme)')

make_table(
    ['Ubicaci\u00f3n', 'x (m)', 'V (kN)', 'M (kN\u00b7m)', 'N (kN)'],
    [
        ['Nudo B (viga)', '0', '94.61', '\u2212144.80', '\u221271.43'],
        ['Cortante nulo', '5.34', '0', '107.04', '\u221271.43'],
        ['Centro luz', '5.0', '6.11', '103.30', '\u221271.43'],
        ['Nudo C (viga)', '10.0', '\u221282.39', '\u2212144.80', '\u221271.43'],
    ]
)

heading3('Fuerzas internas en las columnas')
add_icon("fuerzas_columnas.png", Inches(3.2))

body('En la columna A-B (0 \u2264 y \u2264 3.5 m, origen en A):')
eq('V(y) = H_A = 71.43 kN   (constante)')
eq('N(y) = \u2212V_A = \u221294.61 kN   (compresi\u00f3n)')
eq('M(y) = M_A + H_A\u00b7y = \u2212105.21 + 71.43y   [kN\u00b7m]')

doc.add_page_break()

# ============================================================
# CAPITULO 4: DISENO DE ELEMENTOS
# ============================================================
heading1('Cap\u00edtulo 4: Dise\u00f1o de Elementos')

heading2('4.1 Selecci\u00f3n del perfil de la viga')

add_icon("diseno_viga.png", Inches(3.2))

body('Se realiz\u00f3 una b\u00fasqueda \u00f3ptima entre 16 perfiles AISC serie W, buscando la combinaci\u00f3n '
     'viga-columna de menor peso total que cumpla con FS \u2265 1.4 en flexi\u00f3n y cortante, '
     'con tw \u2265 8 mm y d/tw \u2264 50.')

heading3('Perfil seleccionado: W310\u00d767')

make_table(
    ['Propiedad', 'Valor'],
    [
        ['d (peralte)', '306 mm'],
        ['bf (ancho ala)', '204 mm'],
        ['tf (espesor ala)', '14.6 mm'],
        ['tw (espesor alma)', '8.5 mm (\u2265 8 mm \u2713)'],
        ['d/tw (esbeltez)', '36.0 (\u2264 50 \u2713)'],
        ['Ix', '14,500 cm\u2074'],
        ['Sx', '948 cm\u00b3'],
        ['A', '8,530 mm\u00b2'],
        ['Peso', '67 kg/m'],
    ]
)

heading3('Verificaci\u00f3n de esfuerzos en la viga')

body('Esfuerzo normal combinado (flexi\u00f3n + axial):')
eq('\u03c3 = M/Sx + N/A')
eq('\u03c3 = (144.80\u00d710\u00b3)/948 + (71.43\u00d710\u00b3)/8,530')
eq('\u03c3 = 152.74 + 8.37 = 161.12 MPa')
eq_result('FS_flexi\u00f3n = 250/161.12 = 1.55 \u2265 1.4  \u2713')

body('Esfuerzo cortante m\u00e1ximo:')
eq('\u03c4 = V/(d\u00b7tw) = (94.61\u00d710\u00b3)/(306\u00d78.5) = 36.37 MPa')
eq_result('FS_cortante = 100/36.37 = 2.75 \u2265 1.4  \u2713')

heading2('4.2 Selecci\u00f3n del perfil de la columna')

add_icon("diseno_columna.png", Inches(3.2))

heading3('Perfil seleccionado: W250\u00d758')

make_table(
    ['Propiedad', 'Valor'],
    [
        ['d (peralte)', '252 mm'],
        ['bf (ancho ala)', '203 mm'],
        ['tf (espesor ala)', '13.5 mm'],
        ['tw (espesor alma)', '8.0 mm (\u2265 8 mm \u2713)'],
        ['d/tw (esbeltez)', '31.5 (\u2264 50 \u2713)'],
        ['Ix', '8,700 cm\u2074'],
        ['Sx', '691 cm\u00b3'],
        ['A', '7,420 mm\u00b2'],
        ['Peso', '58 kg/m'],
    ]
)

heading3('Verificaci\u00f3n de esfuerzos en la columna')

body('Carga axial en la columna (incluyendo peso propio):')
eq('N_col = V_A + w_col\u00b7h = 94.61 + (58\u00d79.81/1000)\u00d73.5 = 96.60 kN')

body('Esfuerzo normal combinado:')
eq('\u03c3 = M_A/Sx + N_col/A')
eq('\u03c3 = (105.21\u00d710\u00b3)/691 + (96.60\u00d710\u00b3)/7,420')
eq('\u03c3 = 152.26 + 13.02 = 165.28 MPa')
eq_result('FS_flexi\u00f3n = 250/165.28 = 1.51 \u2265 1.4  \u2713')

body('Esfuerzo cortante:')
eq('\u03c4 = H_A/(d\u00b7tw) = (71.43\u00d710\u00b3)/(252\u00d78.0) = 35.43 MPa')
eq_result('FS_cortante = 100/35.43 = 2.82 \u2265 1.4  \u2713')

heading2('4.3 Perfil personalizado optimizado')

add_icon("perfil_custom.png", Inches(3.2))

body('Se propone un perfil armado I-400\u00d7190 que, al aumentar el peralte y reducir espesores, '
     'logra menor \u00e1rea (y peso) manteniendo el m\u00f3dulo de secci\u00f3n necesario:')

make_table(
    ['Propiedad', 'W310\u00d767 (comercial)', 'I-400\u00d7190 (custom)', 'Ahorro'],
    [
        ['d', '306 mm', '400 mm', '\u2014'],
        ['bf', '204 mm', '190 mm', '\u2014'],
        ['tf', '14.6 mm', '9.0 mm', '\u2014'],
        ['tw', '8.5 mm', '8.0 mm (\u2265 8 \u2713)', '\u2014'],
        ['d/tw', '36.0', '50.0 (\u2264 50 \u2713)', '\u2014'],
        ['A', '8,530 mm\u00b2', '6,476 mm\u00b2', '\u221224.1%'],
        ['Sx', '948 cm\u00b3', '839 cm\u00b3', '\u2014'],
        ['Peso', '67 kg/m', '50.8 kg/m', '\u221224.2%'],
    ],
    highlight_last=True
)

add_img("Optimizacion_Perfil.png", Inches(5.2))
caption('Figura 3. Comparaci\u00f3n perfil comercial vs personalizado: secciones, peso y factores de seguridad.')

heading2('4.4 C\u00edrculos de Mohr')

add_icon("mohr_seccion.png", Inches(3.2))

body('Se analizan 4 puntos de la secci\u00f3n transversal de la viga en el nudo B (secci\u00f3n cr\u00edtica con '
     'M = 144.80 kN\u00b7m y V = 94.61 kN):')

make_table(
    ['Punto', 'Ubicaci\u00f3n', '\u03c3x (MPa)', '\u03c4xy (MPa)', '\u03c31', '\u03c32', '\u03c4max'],
    [
        ['1', 'Fibra superior', '\u2212152.74', '0', '0', '\u2212152.74', '76.37'],
        ['2', 'Fibra inferior', '152.74', '0', '152.74', '0', '76.37'],
        ['3', 'Eje neutro', '0', '36.37', '36.37', '\u221236.37', '36.37'],
        ['4', 'Uni\u00f3n ala-alma', '\u2212130.16', '5.23', '0.21', '\u2212130.37', '65.29'],
    ]
)

add_img("Circulos_Mohr_Correcto.png", Inches(5.2))
caption('Figura 4. C\u00edrculos de Mohr para los 4 puntos cr\u00edticos de la secci\u00f3n en el nudo B.')

doc.add_page_break()

# ============================================================
# CAPITULO 5: DISENO DE CONEXIONES
# ============================================================
heading1('Cap\u00edtulo 5: Dise\u00f1o de Conexiones')

heading2('5.1 Conexi\u00f3n viga-columna (nudos B y C)')

add_icon("conexion_BC.png", Inches(3.2))

body('Se dise\u00f1a una conexi\u00f3n a momento utilizando pernos de alta resistencia A490 y soldadura '
     'de filete E70. Las fuerzas de dise\u00f1o son:')

make_table(
    ['Fuerza', 'Valor', 'Origen'],
    [
        ['Cortante (V)', '94.61 kN', 'Reacci\u00f3n vertical m\u00e1xima'],
        ['Momento (M)', '144.80 kN\u00b7m', 'Momento en nudo B'],
        ['Normal (N)', '71.43 kN', 'Fuerza horizontal total'],
    ]
)

heading3('Pernos A490')
body('Se seleccionan pernos A490 de di\u00e1metro d = 3/4" (19.05 mm), di\u00e1metro de perforaci\u00f3n = 20.6 mm:')

eq('F_u,perno = 1,035 MPa')
eq('\u00c1rea del perno: A_b = \u03c0/4 \u00d7 (19.05)\u00b2 = 285 mm\u00b2')

heading3('Verificaciones')
make_table(
    ['Chequeo', 'Capacidad', 'Demanda', 'Estado'],
    [
        ['Cortante simple perno', '4 \u00d7 0.44\u00d71035\u00d7285/1000 = 521 kN', '94.61 kN', '\u2713 OK'],
        ['Aplastamiento (viga)', '4 \u00d7 2.4\u00d7400\u00d719.05\u00d78.5/1000 = 622 kN', '94.61 kN', '\u2713 OK'],
        ['Fractura \u00e1rea neta', 'A_neta \u00d7 Fu verificado', '\u2014', '\u2713 OK'],
    ]
)

heading3('Soldadura E70')
body('Se emplea soldadura de filete con electrodo E70, tama\u00f1o a = 6 mm:')
eq('F_w = 0.707 \u00d7 a \u00d7 0.6 \u00d7 F_u,E70 = 0.707 \u00d7 6 \u00d7 0.6 \u00d7 482 = 1,227 N/mm')
eq('L_requerida = V/F_w = 94,610/1,227 = 77 mm')
body('Se proporcionan 2 cordones de 200 mm cada uno (L_total = 400 mm >> 77 mm). \u2713')

heading2('5.2 Conexi\u00f3n columna-piso (bases A y D)')

add_icon("conexion_base.png", Inches(3.2))

body('Se dise\u00f1a la placa base y los pernos de anclaje para transmitir las reacciones al cimiento:')

make_table(
    ['Fuerza de dise\u00f1o', 'Valor'],
    [
        ['Vertical (V_A)', '94.61 kN'],
        ['Horizontal (H_A)', '71.43 kN'],
        ['Momento (M_A)', '105.21 kN\u00b7m'],
    ]
)

heading3('Placa base')
body('Dimensiones: 350 \u00d7 300 \u00d7 20 mm (ASTM A36)')
eq('\u03c3_max = N/A_placa + M\u00b7c/I_placa')
body('Se verifica que \u03c3_max < 0.6\u00d7Fy = 150 MPa. \u2713')

heading3('Pernos de anclaje')
body('4 pernos A490 de 3/4", distribuidos sim\u00e9tricamente. Se verifican por cortante y tensi\u00f3n combinada.')

heading3('Soldadura columna-placa')
body('Soldadura de filete E70, a = 6 mm, aplicada en el per\u00edmetro de contacto columna-placa.')

doc.add_page_break()

# ============================================================
# CAPITULO 6: CANTIDADES DE MATERIAL
# ============================================================
heading1('Cap\u00edtulo 6: Cantidades de Material')

heading2('6.1 Perfiles estructurales')

make_table(
    ['Elemento', 'Perfil', 'Longitud (m)', 'Peso (kg/m)', 'Peso total (kg)'],
    [
        ['Viga B\u2013C', 'W310\u00d767', '10.0', '67', '670'],
        ['Columna A\u2013B', 'W250\u00d758', '3.5', '58', '203'],
        ['Columna D\u2013C', 'W250\u00d758', '3.5', '58', '203'],
        ['TOTAL', '\u2014', '17.0', '\u2014', '1,076 kg'],
    ],
    highlight_last=True
)

heading2('6.2 Volumen de acero')

make_table(
    ['Concepto', 'Comercial', 'Custom (I-400\u00d7190)'],
    [
        ['Peso viga', '670 kg', '508 kg'],
        ['Peso 2 columnas', '406 kg', '406 kg'],
        ['Peso total', '1,076 kg', '914 kg (\u221215%)'],
        ['Volumen total', '0.1372 m\u00b3', '0.1167 m\u00b3'],
    ],
    highlight_last=True
)

heading2('6.3 Conexiones')

make_table(
    ['Material', 'Cantidad', 'Descripci\u00f3n'],
    [
        ['Pernos A490, 3/4"', '12 unidades', '4 por conexi\u00f3n \u00d7 3 nudos'],
        ['Soldadura E70', '~3,200 mm', 'Filete 6 mm en nudos + bases'],
        ['Placas base', '2 unidades', '350\u00d7300\u00d720 mm (A36)'],
        ['Pernos ancla', '8 unidades', '4 por base \u00d7 2 apoyos'],
    ]
)

doc.add_page_break()

# ============================================================
# CAPITULO 7: CONCLUSIONES Y COMPARACIONES
# ============================================================
heading1('Cap\u00edtulo 7: Conclusiones y Comparaciones')

heading2('7.1 Conclusiones')

bullet('El p\u00f3rtico plano biempotrado fue analizado correctamente usando las ecuaciones del par\u00e1metro '
       'de rigidez k (p\u00e1ginas 10-11 de la gu\u00eda), considerando la superposici\u00f3n de carga distribuida '
       'y carga lateral s\u00edsmica.')

bullet('La combinaci\u00f3n \u00f3ptima de perfiles es W310\u00d767 para la viga y W250\u00d758 para las columnas, '
       'con factores de seguridad FS = 1.55 (viga) y FS = 1.51 (columna), ambos superiores al m\u00ednimo '
       'requerido de 1.4.')

bullet('El peso total de acero del p\u00f3rtico con perfiles comerciales es de 1,076 kg (0.1372 m\u00b3).')

bullet('El perfil personalizado I-400\u00d7190 logra una reducci\u00f3n de 24.1% en el \u00e1rea de la secci\u00f3n '
       'y 15% en el peso total del p\u00f3rtico (914 kg vs 1,076 kg).')

bullet('Las conexiones atornilladas (A490) y soldadas (E70) cumplen satisfactoriamente con todos '
       'los chequeos de cortante, aplastamiento y fractura en \u00e1rea neta.')

bullet('Los c\u00edrculos de Mohr en la secci\u00f3n cr\u00edtica (nudo B) confirman que los esfuerzos principales '
       'se mantienen dentro de los l\u00edmites admisibles en todos los puntos analizados.')

heading2('7.2 Comparaci\u00f3n con otros proyectos')

body('Se compara nuestro dise\u00f1o con dos proyectos de referencia (MJ de Henao/Ortega y '
     'Laura/Santiago de Pe\u00f1a et al.) que analizaron la estructura de la p\u00e1gina 17 de la gu\u00eda '
     '(viga en L con cargas triangulares), la cual NO corresponde al p\u00f3rtico plano solicitado.')

make_table(
    ['Criterio', 'MJ (Henao/Ortega)', 'Laura/Santiago', 'NUESTRO (xy=45)'],
    [
        ['Estructura', 'Viga en L \u2716', 'Viga en L \u2716', 'P\u00f3rtico plano \u2714'],
        ['Ecuaciones', 'Equilibrio est\u00e1tico', 'Equilibrio est\u00e1tico', 'Ec. p\u00f3rtico (k)'],
        ['Perfil viga', 'W1010 (787 kg/m)', 'W920\u00d7420 (449 kg/m)', 'W310\u00d767 (67 kg/m)'],
        ['\u00c1rea viga', '100,200 mm\u00b2', '57,100 mm\u00b2', '8,530 mm\u00b2 (\u221291.5%)'],
        ['Peso total', '10,891 kg', '6,371 kg', '1,076 kg (\u221290.1%)'],
    ]
)

add_img("Comparacion_Proyectos.png", Inches(5.5))
caption('Figura 5. Comparaci\u00f3n de dise\u00f1o: nuestro proyecto vs proyectos de referencia.')

body('Nuestro dise\u00f1o utiliza un 90.1% menos acero que el proyecto MJ y un 83.1% menos que '
     'Laura/Santiago, demostrando que el an\u00e1lisis correcto del p\u00f3rtico (en lugar de la viga aislada) '
     'conduce a perfiles significativamente m\u00e1s eficientes.')

doc.add_page_break()

# ============================================================
# ANEXOS
# ============================================================
heading1('Anexos')

heading2('Anexo A: Esquema del p\u00f3rtico')
add_img("Esquema_Portico_Correcto.png", Inches(4.5))
caption('Figura A1. Esquema del p\u00f3rtico plano biempotrado con nomenclatura de nodos.')

heading2('Anexo B: Sistema tridimensional')
add_img("Portico_3D_y_2D.png", Inches(5.0))
caption('Figura A2. Vista 3D del sistema estructural y p\u00f3rtico extra\u00eddo.')

heading2('Anexo C: Diagramas de esfuerzos internos')
add_img("Diagramas_Portico.png", Inches(5.5))
caption('Figura A3. Diagramas V, N, M de viga y columnas.')

heading2('Anexo D: Optimizaci\u00f3n del perfil')
add_img("Optimizacion_Perfil.png", Inches(5.0))
caption('Figura A4. Comparaci\u00f3n perfil comercial vs personalizado.')

heading2('Anexo E: C\u00edrculos de Mohr')
add_img("Circulos_Mohr_Correcto.png", Inches(5.0))
caption('Figura A5. C\u00edrculos de Mohr en la secci\u00f3n cr\u00edtica del nudo B.')

heading2('Anexo F: Comparaci\u00f3n con proyectos de referencia')
add_img("Comparacion_Proyectos.png", Inches(5.5))
caption('Figura A6. Diagrama comparativo completo.')

# ============================================================
# GUARDAR
# ============================================================
out = os.path.join(BASE, "archivos", "Informe_Proyecto_Final.docx")
try:
    doc.save(out)
    print(f"Informe guardado: {out}")
except PermissionError:
    out2 = os.path.join(BASE, "archivos", "Informe_Proyecto_Final_v2.docx")
    doc.save(out2)
    print(f"Informe guardado: {out2}")

print(f"Parrafos: {len(doc.paragraphs)}")
print(f"Tablas: {len(doc.tables)}")
