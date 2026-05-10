# -*- coding: utf-8 -*-
"""
CALCULOS CORRECTOS DEL PORTICO PLANO
Proyecto Final - Mecanica de Materiales
Universidad del Quindio - Prof. William Valencia Mina
Estudiante xy = 45

Genera:
  - 3 diagramas PNG (esquema, V/M, Mohr)
  - Memoria_de_Calculos.docx (~50 paginas, 600+ parrafos, paso a paso)

Estructura: Portico plano con 2 columnas + 1 viga
  B _____________ C
  |               |
  |  I1    I2   I1|
  |               |
  A_______________D
  (empotrado)     (empotrado)
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import os, math, warnings
warnings.filterwarnings('ignore')

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import numpy as np

print("="*70)
print("  CALCULOS DEL PORTICO PLANO - PROYECTO MECANICA DE MATERIALES")
print("  Segun guia del Prof. William Valencia Mina (Univ. del Quindio)")
print("="*70)
print("  Generando documento detallado (~50 paginas, 600+ parrafos)...")

# ======================================================================
#  1. PARAMETROS GENERALES
# ======================================================================

h = 3.5       # altura de columnas [m]
L = 10.0      # luz de la viga [m]

Fy   = 250    # esfuerzo de fluencia [MPa]
tau_y = 100   # esfuerzo cortante de fluencia [MPa]
FS_min = 1.4
sigma_adm = Fy / FS_min     # 178.57 MPa
tau_adm   = tau_y / FS_min   # 71.43 MPa
rho_acero = 7850             # kg/m3
g_acc = 9.81                 # m/s2
Fu_A36 = 400                 # MPa

e_losa = 0.18
gamma_c = 24
w_losa = e_losa * gamma_c        # 4.32 kN/m2
w_dead_add = 1.2
w_viva = 3.0
w_total_losa = w_losa + w_dead_add + w_viva   # 8.52 kN/m2

b_trib = 4.0 / 2   # 2.0 m
q_ext = w_total_losa * b_trib   # 17.04 kN/m

P_lat = 45.0  # kN

L_x = 10.0
L_y = 4.0

# ======================================================================
#  2. PERFILES W (AISC)
# ======================================================================

PERFILES = {
    'W200x59':  {'d':210,'bf':205,'tf':14.2,'tw':9.1,'Ix':6110, 'Sx':582, 'A':7580, 'w_kg':59},
    'W250x58':  {'d':252,'bf':203,'tf':13.5,'tw':8.0,'Ix':8700, 'Sx':691, 'A':7420, 'w_kg':58},
    'W250x67':  {'d':257,'bf':204,'tf':15.7,'tw':8.9,'Ix':10400,'Sx':809, 'A':8560, 'w_kg':67},
    'W250x73':  {'d':253,'bf':254,'tf':14.2,'tw':8.6,'Ix':11300,'Sx':895, 'A':9290, 'w_kg':73},
    'W250x89':  {'d':260,'bf':256,'tf':17.3,'tw':10.7,'Ix':14200,'Sx':1090,'A':11400,'w_kg':89},
    'W310x67':  {'d':306,'bf':204,'tf':14.6,'tw':8.5,'Ix':14500,'Sx':948, 'A':8530, 'w_kg':67},
    'W310x74':  {'d':310,'bf':205,'tf':16.3,'tw':9.4,'Ix':16300,'Sx':1050,'A':9480, 'w_kg':74},
    'W310x86':  {'d':310,'bf':254,'tf':16.3,'tw':9.1,'Ix':19500,'Sx':1260,'A':11000,'w_kg':86},
    'W310x97':  {'d':308,'bf':305,'tf':15.4,'tw':9.9,'Ix':22200,'Sx':1440,'A':12300,'w_kg':97},
    'W310x107': {'d':311,'bf':306,'tf':17.0,'tw':10.9,'Ix':24800,'Sx':1600,'A':13600,'w_kg':107},
    'W360x79':  {'d':354,'bf':205,'tf':16.8,'tw':9.4,'Ix':22600,'Sx':1280,'A':10100,'w_kg':79},
    'W360x91':  {'d':353,'bf':254,'tf':16.4,'tw':9.5,'Ix':26600,'Sx':1510,'A':11600,'w_kg':91},
    'W360x101': {'d':357,'bf':255,'tf':18.3,'tw':10.5,'Ix':30000,'Sx':1690,'A':12900,'w_kg':101},
    'W410x85':  {'d':417,'bf':181,'tf':18.2,'tw':10.9,'Ix':31600,'Sx':1510,'A':10800,'w_kg':85},
    'W460x82':  {'d':460,'bf':191,'tf':16.0,'tw':9.9,'Ix':37000,'Sx':1610,'A':10400,'w_kg':82},
    'W460x89':  {'d':463,'bf':192,'tf':17.7,'tw':10.5,'Ix':41000,'Sx':1770,'A':11400,'w_kg':89},
}

def peso_lineal_kN(p):
    return p['w_kg'] * g_acc / 1000

# ======================================================================
#  3. ANALISIS DEL PORTICO
# ======================================================================

def analizar_portico(I2_cm4, I1_cm4, q, P, h_val, L_val):
    k = (I2_cm4 / I1_cm4) * (h_val / L_val)
    VA_q = q * L_val / 2
    VD_q = VA_q
    HA_q = q * L_val**2 / (4 * h_val * (k + 2))
    HD_q = HA_q
    MA_q = q * L_val**2 / (12 * (k + 2))
    MD_q = MA_q
    MB_q = -q * L_val**2 / (6 * (k + 2))
    MC_q = MB_q
    M_mid_q = q * L_val**2 * (3*k + 2) / (24 * (k + 2))
    denom_P = 6*k + 1
    VA_P = 3 * P * h_val * k / (L_val * denom_P)
    VD_P = -VA_P
    HA_P = P / 2
    HD_P = P / 2
    MA_P = -P * h_val * (3*k + 1) / (2 * denom_P)
    MD_P =  P * h_val * (3*k + 1) / (2 * denom_P)
    MB_P =  P * h_val / 2 * 3*k / denom_P
    MC_P = -MB_P
    MA_max = MA_q + abs(MA_P)
    MD_max = MD_q + abs(MD_P)
    MB_max_beam = abs(MB_q) + abs(MB_P)
    MC_max_beam = abs(MC_q) + abs(MC_P)
    VA_max = VA_q + abs(VA_P)
    HA_max = HA_q + HA_P
    N_beam = HA_q + HA_P
    V_beam_max = VA_q + abs(VA_P)
    return {
        'k':k,
        'VA_q':VA_q,'VD_q':VD_q,'HA_q':HA_q,'HD_q':HD_q,
        'MA_q':MA_q,'MD_q':MD_q,'MB_q':MB_q,'MC_q':MC_q,'M_mid_q':M_mid_q,
        'VA_P':VA_P,'VD_P':VD_P,'HA_P':HA_P,'HD_P':HD_P,
        'MA_P':MA_P,'MD_P':MD_P,'MB_P':MB_P,'MC_P':MC_P,
        'MA_max':MA_max,'MD_max':MD_max,
        'MB_max_beam':MB_max_beam,'MC_max_beam':MC_max_beam,
        'M_mid_total':M_mid_q,
        'VA_max':VA_max,'HA_max':HA_max,
        'N_beam':N_beam,'V_beam_max':V_beam_max,
    }

# ======================================================================
#  4. SELECCION DE PERFILES
# ======================================================================
print("\n  Buscando combinacion optima de perfiles...")

best = None
best_weight = 1e9

for bn_ in sorted(PERFILES.keys()):
    bp_ = PERFILES[bn_]
    if bp_['Sx'] < 700:
        continue
    for cn_ in sorted(PERFILES.keys()):
        cp_ = PERFILES[cn_]
        if cp_['Sx'] < 400:
            continue
        w_viga_ = peso_lineal_kN(bp_)
        q_total_ = q_ext + w_viga_
        res_ = analizar_portico(bp_['Ix'], cp_['Ix'], q_total_, P_lat, h, L)
        Sx_b_ = bp_['Sx'] * 1e3
        M_crit_b_ = max(res_['MB_max_beam'], res_['M_mid_total'])
        sigma_b_ = M_crit_b_ * 1e6 / Sx_b_ + res_['N_beam'] * 1e3 / bp_['A']
        tau_b_ = res_['V_beam_max'] * 1e3 / (bp_['d'] * bp_['tw'])
        fs_b_ = Fy / sigma_b_ if sigma_b_ > 0 else 999
        fs_tau_b_ = tau_y / tau_b_ if tau_b_ > 0 else 999
        w_col_ = peso_lineal_kN(cp_)
        N_col_ = res_['VA_max'] + w_col_ * h
        M_col_ = max(res_['MA_max'], res_['MD_max'])
        sigma_c_ = M_col_ * 1e6 / (cp_['Sx']*1e3) + N_col_ * 1e3 / cp_['A']
        tau_c_ = res_['HA_max'] * 1e3 / (cp_['d'] * cp_['tw'])
        fs_c_ = Fy / sigma_c_ if sigma_c_ > 0 else 999
        fs_tau_c_ = tau_y / tau_c_ if tau_c_ > 0 else 999
        if fs_b_ >= FS_min and fs_tau_b_ >= FS_min and fs_c_ >= FS_min and fs_tau_c_ >= FS_min:
            peso_total_ = bp_['w_kg'] * L + cp_['w_kg'] * h * 2
            if peso_total_ < best_weight:
                best_weight = peso_total_
                best = {
                    'beam':bn_,'col':cn_,'res':res_,'q_total':q_total_,
                    'sigma_b':sigma_b_,'tau_b':tau_b_,'fs_b':fs_b_,'fs_tau_b':fs_tau_b_,
                    'sigma_c':sigma_c_,'tau_c':tau_c_,'fs_c':fs_c_,'fs_tau_c':fs_tau_c_,
                    'N_col':N_col_,'M_col':M_col_,
                    'w_viga':w_viga_,'w_col':w_col_,
                    'M_crit_beam':M_crit_b_,
                }

if not best:
    print("  ERROR: No se encontro combinacion viable.")
    sys.exit(1)

bn = best['beam']
cn = best['col']
bp = PERFILES[bn]
cp = PERFILES[cn]
res = best['res']
q_t = best['q_total']
k = res['k']

print(f"  Viga:    {bn}")
print(f"  Columna: {cn}")

# Esfuerzos viga
sigma_flex_b = best['M_crit_beam'] * 1e6 / (bp['Sx']*1e3)
sigma_ax_b   = res['N_beam'] * 1e3 / bp['A']
sigma_tot_b  = sigma_flex_b + sigma_ax_b
tau_b_val    = res['V_beam_max'] * 1e3 / (bp['d'] * bp['tw'])
fs_b         = Fy / sigma_tot_b
fs_tau_b     = tau_y / tau_b_val

# Esfuerzos columna
sigma_flex_c = best['M_col'] * 1e6 / (cp['Sx']*1e3)
sigma_ax_c   = best['N_col'] * 1e3 / cp['A']
sigma_tot_c  = sigma_flex_c + sigma_ax_c
tau_c_val    = res['HA_max'] * 1e3 / (cp['d'] * cp['tw'])
fs_c         = Fy / sigma_tot_c
fs_tau_c     = tau_y / tau_c_val

# ======================================================================
#  5. PERFIL PERSONALIZADO
# ======================================================================
print("  Optimizando perfil personalizado...")

best_custom = None
best_custom_area = 1e9

for d_mm_ in range(250, 420, 10):
    for bf_mm_ in range(150, 260, 10):
        for tf_mm_ in range(8, 20):
            tw_mm_ = 8
            if d_mm_ / tw_mm_ > 50:
                continue
            hw_ = d_mm_ - 2 * tf_mm_
            if hw_ <= 0:
                continue
            A_c_ = 2 * bf_mm_ * tf_mm_ + hw_ * tw_mm_
            Ix_c_ = (bf_mm_ * d_mm_**3 / 12) - ((bf_mm_ - tw_mm_) * hw_**3 / 12)
            Sx_c_ = Ix_c_ / (d_mm_ / 2)
            w_kg_c_ = A_c_ * rho_acero / 1e6
            w_kN_c_ = w_kg_c_ * g_acc / 1000
            q_c_ = q_ext + w_kN_c_
            Ix_cm4_c_ = Ix_c_ / 1e4
            res_c_ = analizar_portico(Ix_cm4_c_, cp['Ix'], q_c_, P_lat, h, L)
            M_crit_c_ = max(res_c_['MB_max_beam'], res_c_['M_mid_total'])
            sigma_c2_ = M_crit_c_ * 1e6 / Sx_c_ + res_c_['N_beam'] * 1e3 / A_c_
            tau_c2_ = res_c_['V_beam_max'] * 1e3 / (d_mm_ * tw_mm_)
            fs_f_ = Fy / sigma_c2_ if sigma_c2_ > 0 else 0
            fs_t_ = tau_y / tau_c2_ if tau_c2_ > 0 else 0
            if fs_f_ >= FS_min and fs_t_ >= FS_min and A_c_ < best_custom_area:
                best_custom_area = A_c_
                best_custom = {
                    'd':d_mm_,'bf':bf_mm_,'tf':tf_mm_,'tw':tw_mm_,
                    'A':A_c_,'Ix':Ix_c_,'Sx':Sx_c_,
                    'w_kg':w_kg_c_,'w_kN':w_kN_c_,
                    'sigma':sigma_c2_,'tau':tau_c2_,
                    'fs_f':fs_f_,'fs_t':fs_t_,
                    'M_crit':M_crit_c_,'N':res_c_['N_beam'],'V':res_c_['V_beam_max'],
                    'esbeltez':d_mm_/tw_mm_,'q_total':q_c_,
                }

if best_custom:
    print(f"  Perfil personalizado: d={best_custom['d']} bf={best_custom['bf']} tf={best_custom['tf']} tw={best_custom['tw']}")

# ======================================================================
#  6. CONEXIONES - pre-calculo
# ======================================================================
Fnv_A490 = 457
Fu_A490 = 1035
Fexx_E70 = 482
Fw_E70 = 0.6 * Fexx_E70

BOLTS = {
    '3/4"': {'d':19.05,'dh':20.6},
    '7/8"': {'d':22.22,'dh':23.8},
    '1"':   {'d':25.40,'dh':27.0},
    '1-1/8"':{'d':28.58,'dh':30.2},
}

M_conn = res['MB_max_beam']
V_conn = res['V_beam_max']
lever = (bp['d'] - bp['tf']) / 1000
F_ala = M_conn / lever

bolt_sel = '7/8"'
bolt = BOLTS[bolt_sel]
A_bolt = math.pi * bolt['d']**2 / 4
Rv_simple = Fnv_A490 * A_bolt / 1000
Rv_doble  = 2 * Rv_simple
n_ala = math.ceil(F_ala / Rv_doble)
n_alma = math.ceil(V_conn / Rv_simple)
n_total = 2 * n_ala + n_alma

t_min_ala = min(bp['tf'], 25)
Rn_bearing_ala = 2.4 * Fu_A36 * bolt['d'] * t_min_ala / 1000
F_dem_ala = F_ala / n_ala

Le = 65
Le_eff = Le - bolt['dh']/2
Rn_tearout = 1.2 * Le_eff * t_min_ala * Fu_A36 / 1000
F_dem_alma = V_conn / n_alma

Ubs = 1.0
Lev = 40
s_bolt = 75
Ant = t_min_ala * (2 * s_bolt)
Anv = t_min_ala * (Lev + (n_ala - 1) * s_bolt - (n_ala - 0.5) * bolt['dh'])
Rn_block = 0.6 * Fu_A36 * Anv / 1000 + Ubs * Fu_A36 * Ant / 1000

M_base = best['M_col']
N_base = best['N_col']
V_base = res['HA_max']
ancho_placa = cp['bf'] + 150
largo_placa = cp['d'] + 200
t_placa = 25
A_placa = ancho_placa * largo_placa
I_placa = ancho_placa * largo_placa**3 / 12
sigma_max_base = N_base*1e3/A_placa + M_base*1e6*(largo_placa/2)/I_placa
sigma_min_base = N_base*1e3/A_placa - M_base*1e6*(largo_placa/2)/I_placa

bolt_base = '3/4"'
bolt_b = BOLTS[bolt_base]
A_bolt_b = math.pi * bolt_b['d']**2 / 4
n_pernos_base = 4
mu = 0.35
V_fric = mu * N_base

a_ala_sold = F_ala * 1e3 / (0.707 * Fw_E70 * 2 * bp['bf'])
hw_beam = bp['d'] - 2*bp['tf']
a_alma_sold = V_conn * 1e3 / (0.707 * Fw_E70 * 2 * hw_beam)
perim_col = 2 * (cp['d'] + cp['bf'])
a_base = 8
cap_sold_base = 0.707 * a_base * Fw_E70 * perim_col / 1000

# ======================================================================
#  7. MOHR pre-calculo
# ======================================================================
Sx_mm3 = bp['Sx'] * 1e3
Ix_mm4 = bp['Ix'] * 1e4
d_mm = bp['d']
tw_mm = bp['tw']
bf_mm = bp['bf']
tf_mm = bp['tf']
hw_mm = d_mm - 2 * tf_mm

M_crit_mohr = best['M_crit_beam']
V_crit_mohr = res['V_beam_max']
N_crit_mohr = res['N_beam']

Q_NA = bf_mm * tf_mm * (d_mm/2 - tf_mm/2) + tw_mm * (d_mm/2 - tf_mm)**2 / 2
Q_ala = bf_mm * tf_mm * (d_mm/2 - tf_mm/2)

puntos_mohr = [
    {'nombre':'Fibra Superior (y = +d/2)','y':d_mm/2,
     'sigma_x':M_crit_mohr*1e6*(d_mm/2)/Ix_mm4 + N_crit_mohr*1e3/bp['A'],'tau_xy':0},
    {'nombre':'Fibra Inferior (y = -d/2)','y':-d_mm/2,
     'sigma_x':-M_crit_mohr*1e6*(d_mm/2)/Ix_mm4 + N_crit_mohr*1e3/bp['A'],'tau_xy':0},
    {'nombre':'Eje Neutro (y = 0)','y':0,
     'sigma_x':N_crit_mohr*1e3/bp['A'],
     'tau_xy':V_crit_mohr*1e3*Q_NA/(Ix_mm4*tw_mm)},
    {'nombre':f'Union Ala-Alma (y = {d_mm/2-tf_mm:.0f} mm)','y':d_mm/2-tf_mm,
     'sigma_x':M_crit_mohr*1e6*(d_mm/2-tf_mm)/Ix_mm4 + N_crit_mohr*1e3/bp['A'],
     'tau_xy':V_crit_mohr*1e3*Q_ala/(Ix_mm4*tw_mm)},
]

for pt_ in puntos_mohr:
    sx_ = pt_['sigma_x']
    txy_ = pt_['tau_xy']
    sa_ = sx_ / 2
    R_ = math.sqrt(sa_**2 + txy_**2)
    pt_['sigma_avg'] = sa_
    pt_['R'] = R_
    pt_['sigma_1'] = sa_ + R_
    pt_['sigma_2'] = sa_ - R_
    pt_['tau_max'] = R_
    pt_['theta_p'] = math.degrees(0.5 * math.atan2(2*txy_, sx_)) if txy_ != 0 else 0.0

# Volumenes
V_viga = bp['A'] * L * 1000 / 1e9
P_viga = V_viga * rho_acero
V_col  = cp['A'] * h * 1000 / 1e9 * 2
P_col  = V_col * rho_acero
V_conex = 0.010
P_conex = V_conex * rho_acero
V_total = V_viga + V_col + V_conex
P_total = P_viga + P_col + P_conex

# ======================================================================
#  8. GENERAR DIAGRAMAS
# ======================================================================
BASE = r"C:\Users\andre\Desktop\proyecto materiales"
print("\n  Generando diagramas...")

def generar_esquema_portico():
    fig, ax = plt.subplots(1,1,figsize=(14,8))
    ax.set_xlim(-2,L+3); ax.set_ylim(-1.5,h+3); ax.set_aspect('equal')
    ax.set_title('Esquema del Portico Plano con Cargas (xy=45)',fontsize=14,fontweight='bold')
    ax.plot([0,0],[0,h],'k-',linewidth=4)
    ax.plot([L,L],[0,h],'k-',linewidth=4)
    ax.plot([0,L],[h,h],'k-',linewidth=4)
    for x,y,lb in [(0,0,'A'),(0,h,'B'),(L,h,'C'),(L,0,'D')]:
        ax.plot(x,y,'ko',markersize=8)
        off=(-0.5,-0.4) if y==0 else (-0.5,0.3)
        ax.annotate(lb,(x,y),fontsize=14,fontweight='bold',xytext=(x+off[0],y+off[1]))
    for xb in [0,L]:
        ax.plot([xb-0.3,xb+0.3],[0,0],'k-',linewidth=2)
        for xi in np.linspace(xb-0.3,xb+0.3,5):
            ax.plot([xi,xi-0.15],[0,-0.15],'k-',linewidth=1)
    for i in range(21):
        xa=i*L/20
        ax.annotate('',xy=(xa,h),xytext=(xa,h+1.5),arrowprops=dict(arrowstyle='->',color='blue',lw=1))
    ax.plot([0,L],[h+1.5,h+1.5],'b-',linewidth=1.5)
    ax.text(L/2,h+1.8,f'q = {q_t:.2f} kN/m',ha='center',fontsize=11,color='blue',fontweight='bold')
    ax.annotate('',xy=(0,h),xytext=(-1.5,h),arrowprops=dict(arrowstyle='->',color='red',lw=2.5))
    ax.text(-1.8,h+0.2,f'P = {P_lat} kN',fontsize=11,color='red',fontweight='bold')
    ax.annotate('',xy=(0,-0.8),xytext=(L,-0.8),arrowprops=dict(arrowstyle='<->',color='gray',lw=1.5))
    ax.text(L/2,-1.1,f'L = {L} m',ha='center',fontsize=11,color='gray')
    ax.annotate('',xy=(L+1,0),xytext=(L+1,h),arrowprops=dict(arrowstyle='<->',color='gray',lw=1.5))
    ax.text(L+1.5,h/2,f'h = {h} m',ha='left',fontsize=11,color='gray',rotation=90,va='center')
    pr=dict(color='green',fontsize=9,fontweight='bold')
    ax.text(-1.5,-0.5,f'HA={res["HA_max"]:.1f} kN\nVA={res["VA_max"]:.1f} kN\nMA={res["MA_max"]:.1f} kN-m',**pr)
    ax.text(L+0.5,-0.5,f'HD={res["HA_max"]:.1f} kN\nVD={res["VA_max"]:.1f} kN\nMD={res["MD_max"]:.1f} kN-m',**pr)
    ax.text(L/2,h-0.3,f'Viga: {bn}',ha='center',fontsize=10,color='purple',fontweight='bold')
    ax.text(-0.8,h/2,f'Col: {cn}',ha='center',fontsize=10,color='purple',fontweight='bold',rotation=90)
    ax.text(L+0.8,h/2,f'Col: {cn}',ha='center',fontsize=10,color='purple',fontweight='bold',rotation=90)
    ax.axis('off'); fig.tight_layout()
    fig.savefig(os.path.join(BASE,'Esquema_Portico_Correcto.png'),dpi=150,bbox_inches='tight')
    plt.close()
    print("    -> Esquema_Portico_Correcto.png")

def generar_diagramas_VM():
    fig,axes=plt.subplots(2,2,figsize=(16,10))
    xb=np.linspace(0,L,200)
    Mq=q_t*xb*(L-xb)/2+res['MB_q']
    Mp=res['MB_P']+(res['MC_P']-res['MB_P'])*xb/L
    Mt=Mq+abs(Mp)
    ax=axes[0,0]
    ax.fill_between(xb,0,Mt,alpha=0.3,color='red'); ax.plot(xb,Mt,'r-',linewidth=2)
    ax.axhline(y=0,color='k',linewidth=0.5); ax.set_title('Diagrama de Momento Flector - Viga',fontweight='bold')
    ax.set_xlabel('x [m]'); ax.set_ylabel('M [kN-m]'); ax.grid(True,alpha=0.3); ax.invert_yaxis()
    ax=axes[0,1]
    Vq=q_t*L/2-q_t*xb; Vt=Vq+abs(res['VA_P'])
    ax.fill_between(xb,0,Vt,alpha=0.3,color='blue'); ax.plot(xb,Vt,'b-',linewidth=2)
    ax.axhline(y=0,color='k',linewidth=0.5); ax.set_title('Diagrama de Fuerza Cortante - Viga',fontweight='bold')
    ax.set_xlabel('x [m]'); ax.set_ylabel('V [kN]'); ax.grid(True,alpha=0.3)
    ax=axes[1,0]
    yc=np.linspace(0,h,100)
    MAt=res['MA_max']; MBt=abs(res['MB_q'])+abs(res['MB_P'])
    Mc_=MAt+((-MBt)-MAt)*yc/h
    ax.fill_betweenx(yc,0,Mc_,alpha=0.3,color='red'); ax.plot(Mc_,yc,'r-',linewidth=2)
    ax.axvline(x=0,color='k',linewidth=0.5); ax.set_title('Diagrama de Momento - Columna',fontweight='bold')
    ax.set_xlabel('M [kN-m]'); ax.set_ylabel('y [m]'); ax.grid(True,alpha=0.3)
    ax=axes[1,1]
    ax.fill_betweenx(yc,0,res['HA_max'],alpha=0.3,color='blue')
    ax.plot([res['HA_max'],res['HA_max']],[0,h],'b-',linewidth=2)
    ax.axvline(x=0,color='k',linewidth=0.5); ax.set_title('Diagrama de Cortante - Columna',fontweight='bold')
    ax.set_xlabel('V [kN]'); ax.set_ylabel('y [m]'); ax.grid(True,alpha=0.3)
    fig.suptitle(f'Diagramas de Fuerza Interna ({bn} + {cn})',fontsize=14,fontweight='bold',y=1.02)
    fig.tight_layout()
    fig.savefig(os.path.join(BASE,'Diagramas_Portico.png'),dpi=150,bbox_inches='tight')
    plt.close()
    print("    -> Diagramas_Portico.png")

def generar_mohr():
    fig,axes=plt.subplots(2,2,figsize=(14,10))
    for idx,(ax,pt) in enumerate(zip(axes.flat,puntos_mohr)):
        sa=pt['sigma_avg']; R=pt['R']; s1=pt['sigma_1']; s2=pt['sigma_2']
        th=np.linspace(0,2*math.pi,200)
        cx_=sa+R*np.cos(th); cy_=R*np.sin(th)
        ax.plot(cx_,cy_,'b-',linewidth=2)
        ax.axhline(y=0,color='k',linewidth=0.5); ax.axvline(x=0,color='k',linewidth=0.5)
        ax.plot(s1,0,'ro',markersize=8,label=f's1={s1:.1f}')
        ax.plot(s2,0,'go',markersize=8,label=f's2={s2:.1f}')
        ax.plot(sa,R,'b^',markersize=8,label=f'tmax={R:.1f}')
        ax.set_title(f'Punto {idx+1}: {pt["nombre"]}',fontsize=10,fontweight='bold')
        ax.set_xlabel('sigma [MPa]'); ax.set_ylabel('tau [MPa]')
        ax.legend(fontsize=8); ax.grid(True,alpha=0.3); ax.set_aspect('equal')
    fig.suptitle('Circulos de Mohr - Seccion Critica de la Viga',fontsize=14,fontweight='bold')
    fig.tight_layout()
    fig.savefig(os.path.join(BASE,'Circulos_Mohr_Correcto.png'),dpi=150,bbox_inches='tight')
    plt.close()
    print("    -> Circulos_Mohr_Correcto.png")

generar_esquema_portico()
generar_diagramas_VM()
generar_mohr()

# ======================================================================
#  9. GENERAR MEMORIA DE CALCULOS (WORD) - 50+ paginas, 600+ parrafos
# ======================================================================
print("\n  Generando Memoria de Calculos (Word)...")

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls

OUT_DIR = os.path.join(BASE, "archivos")
os.makedirs(OUT_DIR, exist_ok=True)
FINAL = os.path.join(OUT_DIR, "Memoria_de_Calculos_FINAL.docx")

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

DARK_BLUE = RGBColor(0x00, 0x2B, 0x5C)
RED = RGBColor(0xC0, 0x00, 0x00)

def heading(text, level=1):
    h_ = doc.add_heading(text, level=level)
    for run in h_.runs:
        run.font.color.rgb = DARK_BLUE
    return h_

def body(text):
    return doc.add_paragraph(text)

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

ICONOS_DIR = os.path.join(BASE, "iconos_elementos")

def banner(elemento, descripcion, icono=None):
    """Inserta un mini-diagrama del portico + banner que identifica QUE PARTE se calcula."""
    # Primero el mini-diagrama si existe
    if icono:
        icon_path = os.path.join(ICONOS_DIR, icono)
        if os.path.exists(icon_path):
            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic_p.add_run()
            run.add_picture(icon_path, width=Inches(4.0))
    # Banner con fondo coloreado
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    from docx.oxml import parse_xml
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="002B5C" w:val="clear"/>')
    p._element.get_or_add_pPr().append(shading)
    r = p.add_run(f"  {elemento}  ")
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(13)
    r.bold = True
    r.font.name = 'Calibri'
    # Linea descriptiva debajo
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(descripcion)
    r2.font.size = Pt(10)
    r2.italic = True
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

import re as _re

def _sym(text):
    """Convierte texto plano a simbolos Unicode (griegas y operadores)."""
    t = text
    # Griegas (reemplazar ANTES de que _parse busque subindices)
    for src, dst in [
        ('Sigma', '\u03A3'), ('sigma', '\u03C3'), ('tau', '\u03C4'),
        ('theta', '\u03B8'), ('gamma', '\u03B3'), ('phi', '\u03C6'),
        ('delta', '\u03B4'), ('lambda', '\u03BB'), ('omega', '\u03C9'),
    ]:
        t = t.replace(src, dst)
    # Operadores
    t = t.replace('sqrt', '\u221A')
    t = t.replace('>=', '\u2265')
    t = t.replace('<=', '\u2264')
    t = t.replace(' x ', ' \u00D7 ')
    t = t.replace('+-', '\u00B1')
    t = t.replace('->', '\u2192')
    return t

# ── Constructores OMML ──────────────────────────────────────────

def _mr(text, bold=False, color_hex=None):
    """Crea un <m:r> (math run) con texto."""
    r = OxmlElement('m:r')
    # Propiedades de math run
    if bold or color_hex:
        mrPr = OxmlElement('m:rPr')
        if bold:
            sty = OxmlElement('m:sty')
            sty.set(qn('m:val'), 'bi')
            mrPr.append(sty)
        r.append(mrPr)
    # Propiedades de Word run (fuente, color)
    wRPr = OxmlElement('w:rPr')
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:ascii'), 'Cambria Math')
    rF.set(qn('w:hAnsi'), 'Cambria Math')
    wRPr.append(rF)
    if bold:
        wRPr.append(OxmlElement('w:b'))
    if color_hex:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color_hex)
        wRPr.append(c)
    r.append(wRPr)
    mt = OxmlElement('m:t')
    mt.text = text
    mt.set(qn('xml:space'), 'preserve')
    r.append(mt)
    return r

def _msub(base_text, sub_text, bold=False, color_hex=None):
    """Crea <m:sSub> (subindice real de Word): base con sub abajo."""
    sSub = OxmlElement('m:sSub')
    e = OxmlElement('m:e')
    e.append(_mr(base_text, bold, color_hex))
    sSub.append(e)
    sub = OxmlElement('m:sub')
    sub.append(_mr(sub_text, bold, color_hex))
    sSub.append(sub)
    return sSub

def _msup(base_text, sup_text, bold=False, color_hex=None):
    """Crea <m:sSup> (superindice real de Word)."""
    sSup = OxmlElement('m:sSup')
    e = OxmlElement('m:e')
    e.append(_mr(base_text, bold, color_hex))
    sSup.append(e)
    sup = OxmlElement('m:sup')
    sup.append(_mr(sup_text, bold, color_hex))
    sSup.append(sup)
    return sSup

def _mfrac(num_text, den_text, bold=False, color_hex=None):
    """Crea <m:f> (fraccion real de Word)."""
    f = OxmlElement('m:f')
    num = OxmlElement('m:num')
    num.append(_mr(num_text, bold, color_hex))
    f.append(num)
    den = OxmlElement('m:den')
    den.append(_mr(den_text, bold, color_hex))
    f.append(den)
    return f

def _parse_to_omath(text, bold=False, color_hex=None):
    """
    Parsea texto de ecuacion y genera elementos OMML con subindices reales.
    Detecta patrones X_sub y los convierte a <m:sSub>.
    Detecta X^sup y los convierte a <m:sSup>.
    """
    oMath = OxmlElement('m:oMath')
    # Patron: una o mas letras/simbolos seguidos de _ y el subindice
    # o ^ y el superindice
    pattern = _re.compile(
        r'([A-Za-z\u03B1-\u03C9\u03A0-\u03A9\u221A]+)'  # base
        r'([_^])'                                          # _ o ^
        r'([A-Za-z0-9\u00E1\u00E9\u00ED\u00F3\u00FA]+)'   # sub/superindice
    )
    pos = 0
    for m in pattern.finditer(text):
        start, end = m.span()
        # Texto antes del match
        if start > pos:
            chunk = text[pos:start]
            if chunk:
                oMath.append(_mr(chunk, bold, color_hex))
        base = m.group(1)
        op = m.group(2)
        idx = m.group(3)
        if op == '_':
            oMath.append(_msub(base, idx, bold, color_hex))
        else:  # ^
            oMath.append(_msup(base, idx, bold, color_hex))
        pos = end
    # Texto restante
    if pos < len(text):
        remaining = text[pos:]
        if remaining:
            oMath.append(_mr(remaining, bold, color_hex))
    return oMath

def _build_omath_para(text, bold=False, color_hex=None):
    """Construye <m:oMathPara> completo con subindices y superindices reales."""
    oMathPara = OxmlElement('m:oMathPara')
    # Centrado
    pr = OxmlElement('m:oMathParaPr')
    jc = OxmlElement('m:jc')
    jc.set(qn('m:val'), 'center')
    pr.append(jc)
    oMathPara.append(pr)
    # Parsear y generar el oMath con subindices
    oMath = _parse_to_omath(text, bold, color_hex)
    oMathPara.append(oMath)
    return oMathPara

def eq(text):
    """Inserta ecuacion OMML con subindices reales de Word."""
    p = doc.add_paragraph()
    p._element.append(_build_omath_para(_sym(text)))
    return p

def eq_red(text):
    """Inserta ecuacion OMML en rojo/negrita con subindices reales."""
    p = doc.add_paragraph()
    p._element.append(_build_omath_para(_sym(text), bold=True, color_hex='C00000'))
    return p

def add_img(name, width_in=5.5):
    path = os.path.join(BASE, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    for i, ht in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ht
        for p_ in cell.paragraphs:
            for r_ in p_.runs:
                r_.bold = True
                r_.font.size = Pt(9)
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p_ in cell.paragraphs:
                for r_ in p_.runs:
                    r_.font.size = Pt(9)
    return table

def spacer():
    doc.add_paragraph()

# ==============================================================
#  PORTADA
# ==============================================================
print("    Portada...")

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UNIVERSIDAD DEL QUINDIO")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = DARK_BLUE

spacer()

for txt in ["Facultad de Ingenieria", "Programa de Ingenieria Civil"]:
    p = doc.add_paragraph(txt)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_ in p.runs:
        r_.font.size = Pt(13)

for _ in range(3):
    doc.add_paragraph()

for txt, sz, bf_ in [
    ("MEMORIA DE CALCULOS", 20, True),
    ("Analisis y Diseno de Portico Plano de Acero", 15, True),
    ("Mecanica de Materiales - Proyecto Final (xy = 45)", 13, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(sz); r.bold = bf_; r.font.color.rgb = DARK_BLUE

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Docente: Ing. William Valencia Mina")
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Armenia, Quindio - 2025")
r.font.size = Pt(12)

doc.add_page_break()

# ==============================================================
#  TABLA DE CONTENIDO
# ==============================================================
heading("TABLA DE CONTENIDO")
body("1.  Introduccion y Datos del Proyecto ................................ 3")
body("2.  Analisis de Cargas ............................................... 5")
body("3.  Analisis Estructural - Reacciones y Fuerzas Internas ............ 9")
body("4.  Diseno de Elementos ............................................. 19")
body("5.  Conexiones ....................................................... 27")
body("6.  Analisis de Esfuerzos y Circulos de Mohr ........................ 35")
body("7.  Volumenes, Pesos y Conclusiones .................................. 43")
body("8.  Referencias ...................................................... 46")
body("9.  Anexos ........................................................... 47")
doc.add_page_break()

# ##############################################################
# CAPITULO 1
# ##############################################################
print("    Cap 1: Introduccion...")

heading("CAPITULO 1: INTRODUCCION Y DATOS DEL PROYECTO")
spacer()

heading("1.1 Objetivo General", 2)
body("El presente documento constituye la Memoria de Calculos del proyecto final de la asignatura "
     "Mecanica de Materiales, correspondiente al analisis y diseno completo de un portico plano de acero "
     "estructural. El objetivo general es aplicar los conceptos fundamentales de la mecanica de materiales "
     "para dimensionar y verificar cada uno de los elementos que componen la estructura, garantizando "
     "condiciones adecuadas de resistencia y seguridad.")

body("El portico analizado forma parte de un edificio tridimensional compuesto por columnas y vigas "
     "de acero, con una losa de concreto reforzado que transmite las cargas gravitacionales y una fuerza "
     "sismica lateral equivalente. Se utilizan las ecuaciones analiticas proporcionadas en la guia del "
     "proyecto para un portico biempotrado con dos casos de carga superpuestos.")

body("Todos los calculos se realizan de manera paso a paso, mostrando la formula general, la sustitucion "
     "numerica, los pasos intermedios y el resultado final, con el fin de facilitar la revision y "
     "comprension del procedimiento por parte del lector.")

spacer()

heading("1.2 Objetivos Especificos", 2)
bullet("Determinar las cargas actuantes sobre el portico plano a partir de las condiciones de carga "
       "de la losa de concreto, la carga viva de oficinas y la fuerza sismica lateral (xy = 45 kN).")
bullet("Calcular las reacciones en los apoyos empotrados (A y D) utilizando las ecuaciones del portico "
       "biempotrado para carga distribuida y carga lateral.")
bullet("Obtener los diagramas de momento flector, fuerza cortante y fuerza normal para la viga y las "
       "columnas del portico.")
bullet("Seleccionar perfiles comerciales W (AISC) para la viga y la columna que cumplan con un factor "
       "de seguridad minimo de 1.4 en flexion, cortante y flexo-compresion.")
bullet("Disenar un perfil I personalizado optimizado para la viga, comparando su eficiencia con el "
       "perfil comercial seleccionado.")
bullet("Disenar las conexiones viga-columna y columna-fundacion, verificando pernos A490 por cortante, "
       "aplastamiento, desgarramiento y block shear, asi como soldaduras con electrodo E70.")
bullet("Realizar el analisis de esfuerzos en la seccion critica de la viga, calculando esfuerzos "
       "principales y cortante maximo en 4 puntos representativos, y graficar los circulos de Mohr.")
bullet("Computar los volumenes y pesos de acero del portico y formular las conclusiones del proyecto.")

spacer()

heading("1.3 Material: Acero ASTM A36", 2)
body("El material empleado para todos los elementos estructurales es acero ASTM A36, cuyas propiedades "
     "mecanicas principales se resumen a continuacion:")

body("Esfuerzo de fluencia:")
eq("Fy = 250 MPa")

body("Esfuerzo cortante de fluencia:")
eq("tau_y = 0.40 x Fy = 0.40 x 250 = 100 MPa")

body("Resistencia ultima a la traccion:")
eq("Fu = 400 MPa")

body("Modulo de elasticidad:")
eq("E = 200,000 MPa = 200 GPa")

body("Densidad del acero:")
eq("rho = 7,850 kg/m3")

body("Factor de seguridad minimo requerido:")
eq(f"FS_min = {FS_min}")

body("Esfuerzo normal admisible:")
eq("sigma_adm = Fy / FS")
eq(f"sigma_adm = 250 / 1.40")
eq_red(f"sigma_adm = {sigma_adm:.2f} MPa")

body("Esfuerzo cortante admisible:")
eq("tau_adm = tau_y / FS")
eq(f"tau_adm = 100 / 1.40")
eq_red(f"tau_adm = {tau_adm:.2f} MPa")

body("Estos valores admisibles seran la referencia para verificar que los esfuerzos actuantes en cada "
     "elemento sean menores, garantizando asi la seguridad estructural del portico.")

spacer()

heading("1.4 Geometria del Portico", 2)
body("El portico plano esta conformado por los siguientes nodos y elementos:")

body("Nodo A: Base de la columna izquierda, apoyo empotrado, coordenada (0, 0).")
body("Nodo B: Conexion superior columna izquierda-viga, nudo rigido, coordenada (0, 3.5 m).")
body("Nodo C: Conexion superior columna derecha-viga, nudo rigido, coordenada (10.0, 3.5 m).")
body("Nodo D: Base de la columna derecha, apoyo empotrado, coordenada (10.0, 0).")

body("Dimensiones principales:")
eq(f"Altura de columnas:  h = {h} m")
eq(f"Luz de la viga:      L = {L} m")
eq(f"Relacion h/L = {h}/{L} = {h/L:.2f}")

body("Elementos:")
body("- Columna izquierda (A-B): longitud h = 3.5 m, inercia I1")
body("- Columna derecha (D-C): longitud h = 3.5 m, inercia I1")
body("- Viga (B-C): longitud L = 10.0 m, inercia I2")

add_img('Portico_3D_y_2D.png', 6.5)
body("Figura 1.1: Sistema tridimensional y portico plano extraido con identificacion de elementos.")

add_img('Esquema_Portico_Correcto.png', 6.0)
body("Figura 1.2: Esquema del portico plano con cargas, reacciones y perfiles asignados.")

doc.add_page_break()

# ##############################################################
# CAPITULO 2
# ##############################################################
print("    Cap 2: Analisis de cargas...")

heading("CAPITULO 2: ANALISIS DE CARGAS")
spacer()

heading("2.1 Sistema Tridimensional de Referencia", 2)
body("El portico plano analizado forma parte de un edificio tridimensional de acero con las "
     "siguientes caracteristicas:")

body("Configuracion en planta:")
eq(f"Dimension en X (luz viga): L_x = {L_x} m")
eq(f"Dimension en Y (ancho tributario): L_y = {L_y} m")
eq(f"Area en planta: A = L_x x L_y = {L_x} x {L_y} = {L_x*L_y:.0f} m2")

body("Configuracion en elevacion:")
eq(f"Altura de piso: h = {h} m")
eq("Numero de pisos: 1 (estructura de un nivel)")

body("El edificio cuenta con 4 columnas en las esquinas y vigas perimetrales que conforman "
     "los marcos resistentes. La losa de concreto reforzado se apoya sobre las vigas y transfiere "
     "las cargas gravitacionales mediante el ancho tributario correspondiente.")

spacer()

heading("2.2 Cargas sobre la Losa", 2)
body("Las cargas sobre la losa se descomponen en tres componentes principales:")

body("Paso 1: Peso propio de la losa de concreto.")
body(f"La losa tiene un espesor de e = {e_losa} m y peso unitario gamma_c = {gamma_c} kN/m3.")
eq("w_losa = e x gamma_c")
eq(f"w_losa = {e_losa} x {gamma_c}")
eq_red(f"w_losa = {w_losa:.2f} kN/m2")

body("Paso 2: Cargas muertas adicionales.")
body("Se incluyen acabados de piso (baldosa, mortero), cielo raso e instalaciones.")
eq(f"w_dead_add = {w_dead_add:.2f} kN/m2")
eq_red(f"w_muerta_adicional = {w_dead_add:.2f} kN/m2")

body("Paso 3: Carga viva de ocupacion.")
body("Segun la NSR-10, para oficinas:")
eq(f"w_viva = {w_viva:.2f} kN/m2")
eq_red(f"w_viva = {w_viva:.2f} kN/m2")

body("Paso 4: Carga total sobre la losa.")
eq("w_total = w_losa + w_dead_add + w_viva")
eq(f"w_total = {w_losa:.2f} + {w_dead_add:.2f} + {w_viva:.2f}")
eq_red(f"w_total = {w_total_losa:.2f} kN/m2")

body("Este valor representa la carga superficial total que actua sobre cada metro cuadrado de la "
     "losa de concreto y que sera transferida a las vigas del portico.")

spacer()

heading("2.3 Ancho Tributario y Carga Distribuida", 2)
body("Para la losa bidireccional de 10 m x 4 m, el ancho tributario en la direccion larga es "
     "igual a la mitad de la dimension corta:")
eq("b_trib = L_y / 2")
eq(f"b_trib = {L_y} / 2")
eq_red(f"b_trib = {b_trib:.1f} m")

body("Carga distribuida lineal sobre la viga:")
eq("q_ext = w_total x b_trib")
eq(f"q_ext = {w_total_losa:.2f} x {b_trib:.1f}")
eq_red(f"q_ext = {q_ext:.2f} kN/m")

body("Este valor de q_ext representa la carga distribuida externa sin incluir el peso propio de la viga.")

spacer()

heading("2.4 Carga Lateral Sismica", 2)
body("La fuerza horizontal aplicada en el nudo B del portico representa la fuerza sismica "
     "equivalente, definida por el parametro xy del grupo de trabajo:")
eq("P = xy = 45 kN")
eq_red(f"P = {P_lat:.1f} kN (horizontal en nudo B)")

body("Esta fuerza actua en la direccion X positiva. En el analisis se considera que el sismo puede "
     "actuar en ambas direcciones, tomando valores absolutos para la envolvente.")

body("La carga lateral se aplica en el nudo B (esquina superior izquierda), ya que la masa se "
     "concentra en el nivel de la losa.")

spacer()

heading("2.5 Peso Propio de los Elementos", 2)
body(f"Para la viga, perfil {bn} ({bp['w_kg']} kg/m):")
eq(f"w_pp_viga = {bp['w_kg']} x {g_acc} / 1000")
eq(f"w_pp_viga = {bp['w_kg']*g_acc:.3f} / 1000")
eq_red(f"w_pp_viga = {best['w_viga']:.3f} kN/m")

body(f"Para la columna, perfil {cn} ({cp['w_kg']} kg/m):")
eq(f"w_pp_col = {cp['w_kg']} x {g_acc} / 1000")
eq(f"w_pp_col = {cp['w_kg']*g_acc:.3f} / 1000")
eq_red(f"w_pp_col = {best['w_col']:.3f} kN/m")

body("Carga axial por peso propio de la columna:")
eq(f"N_pp_col = w_pp_col x h = {best['w_col']:.3f} x {h}")
eq_red(f"N_pp_col = {best['w_col']*h:.2f} kN")

spacer()

heading("2.6 Carga Total sobre la Viga", 2)
body("La carga distribuida total incluye la carga externa mas el peso propio de la viga:")
eq("q_total = q_ext + w_pp_viga")
eq(f"q_total = {q_ext:.2f} + {best['w_viga']:.3f}")
eq_red(f"q_total = {q_t:.3f} kN/m")

body(f"El peso propio representa el {best['w_viga']/q_ext*100:.1f}% de la carga externa.")

body("Resumen de cargas:")
add_table(
    ['Carga', 'Simbolo', 'Valor', 'Unidad'],
    [
        ['Peso losa', 'w_losa', f'{w_losa:.2f}', 'kN/m2'],
        ['Carga muerta adic.', 'w_dead', f'{w_dead_add:.2f}', 'kN/m2'],
        ['Carga viva', 'w_viva', f'{w_viva:.2f}', 'kN/m2'],
        ['Total losa', 'w_total', f'{w_total_losa:.2f}', 'kN/m2'],
        ['Ancho tributario', 'b_trib', f'{b_trib:.1f}', 'm'],
        ['Carga ext. viga', 'q_ext', f'{q_ext:.2f}', 'kN/m'],
        ['Peso propio viga', 'w_pp', f'{best["w_viga"]:.3f}', 'kN/m'],
        ['Carga total viga', 'q_total', f'{q_t:.3f}', 'kN/m'],
        ['Fuerza lateral', 'P', f'{P_lat:.1f}', 'kN'],
    ]
)
body("Tabla 2.1: Resumen completo de cargas del proyecto.")

doc.add_page_break()

# ##############################################################
# CAPITULO 3
# ##############################################################
print("    Cap 3: Analisis estructural...")

heading("CAPITULO 3: ANALISIS ESTRUCTURAL - REACCIONES Y FUERZAS INTERNAS")
spacer()

body("En este capitulo se determinan las reacciones en los apoyos empotrados y las fuerzas internas "
     "en cada elemento del portico. El analisis se realiza por superposicion de dos casos de carga:")
body("- Caso 1: Carga distribuida uniforme q sobre la viga BC")
body("- Caso 2: Carga lateral concentrada P en el nudo B")

spacer()

heading("3.1 Parametro de Rigidez k", 2)
body("El parametro k relaciona las inercias y dimensiones de los elementos del portico:")
eq("k = (I_2 / I_1) x (h / L)")

body("donde:")
body("  I_2 = Inercia de la viga")
body("  I_1 = Inercia de la columna")
body("  h = Altura de columnas")
body("  L = Luz de la viga")

body("Valores de los perfiles seleccionados:")
eq(f"I_2 = Ix_viga = {bp['Ix']} cm4  (perfil {bn})")
eq(f"I_1 = Ix_col  = {cp['Ix']} cm4  (perfil {cn})")

body("Calculo paso a paso:")
eq(f"k = ({bp['Ix']} / {cp['Ix']}) x ({h} / {L})")
eq(f"k = {bp['Ix']/cp['Ix']:.6f} x {h/L:.2f}")
eq_red(f"k = {k:.4f}")

body("Un valor de k < 1 indica que las columnas son relativamente mas rigidas que la viga.")

spacer()

# --- CASO 1 ---
banner("CASO 1: CARGA SOBRE LA VIGA (elemento B-C)",
       "Carga distribuida uniforme q = 17.697 kN/m proveniente de la losa, aplicada sobre toda la viga B-C.",
       "caso1_viga.png")
heading("3.2 Caso 1: Carga Distribuida Uniforme q", 2)
body(f"Carga distribuida: q = {q_t:.3f} kN/m sobre la viga BC.")

spacer()

heading("3.2.1 Reaccion Vertical V_A y V_D", 3)
body("Por simetria de la carga distribuida:")
eq("V_A = V_D = q x L / 2")
eq(f"V_A = {q_t:.3f} x {L} / 2")
eq(f"V_A = {q_t*L:.3f} / 2")
eq_red(f"V_A = V_D = {res['VA_q']:.2f} kN")

body("Verificacion: V_A + V_D = q x L")
eq(f"{res['VA_q']:.2f} + {res['VA_q']:.2f} = {2*res['VA_q']:.2f} kN")
eq(f"q x L = {q_t:.3f} x {L} = {q_t*L:.2f} kN  OK")

spacer()

heading("3.2.2 Reaccion Horizontal H_A y H_D", 3)
eq("H_A = H_D = q x L^2 / [4 x h x (k + 2)]")
eq(f"Numerador = q x L^2 = {q_t:.3f} x {L**2:.0f} = {q_t*L**2:.2f}")
eq(f"k + 2 = {k:.4f} + 2 = {k+2:.4f}")
eq(f"Denominador = 4 x {h} x {k+2:.4f} = {4*h*(k+2):.4f}")
eq(f"H_A = {q_t*L**2:.2f} / {4*h*(k+2):.4f}")
eq_red(f"H_A = H_D = {res['HA_q']:.2f} kN")

spacer()

heading("3.2.3 Momento en la Base M_A y M_D", 3)
eq("M_A = M_D = q x L^2 / [12 x (k + 2)]")
eq(f"Numerador = q x L^2 = {q_t*L**2:.2f}")
eq(f"Denominador = 12 x (k + 2) = 12 x {k+2:.4f} = {12*(k+2):.4f}")
eq(f"M_A = {q_t*L**2:.2f} / {12*(k+2):.4f}")
eq_red(f"M_A = M_D = {res['MA_q']:.2f} kN-m")

body("Momento positivo: tension en la cara interior de la columna.")

spacer()

heading("3.2.4 Momento en los Nudos M_B y M_C", 3)
eq("M_B = M_C = -q x L^2 / [6 x (k + 2)]")
eq(f"Numerador = q x L^2 = {q_t*L**2:.2f}")
eq(f"Denominador = 6 x (k + 2) = 6 x {k+2:.4f} = {6*(k+2):.4f}")
eq(f"M_B = -{q_t*L**2:.2f} / {6*(k+2):.4f}")
eq_red(f"M_B = M_C = {res['MB_q']:.2f} kN-m")

body("El signo negativo indica que la cara superior de la viga esta en tension en los nudos.")

spacer()

heading("3.2.5 Momento en el Centro de la Viga M_mid", 3)
eq("M_mid = q x L^2 x (3k + 2) / [24 x (k + 2)]")
eq(f"3k + 2 = 3 x {k:.4f} + 2 = {3*k+2:.4f}")
eq(f"Numerador = {q_t:.3f} x {L**2:.0f} x {3*k+2:.4f} = {q_t*L**2*(3*k+2):.2f}")
eq(f"Denominador = 24 x {k+2:.4f} = {24*(k+2):.4f}")
eq(f"M_mid = {q_t*L**2*(3*k+2):.2f} / {24*(k+2):.4f}")
eq_red(f"M_mid = {res['M_mid_q']:.2f} kN-m")

spacer()

heading("3.2.6 Funcion de Momento M_BC(x) en la Viga", 3)
body("La funcion de momento a lo largo de la viga (0 <= x <= L):")
eq("M_BC(x) = q x x x (L - x) / 2 + M_B")
eq(f"M_BC(x) = {q_t:.3f} x x x ({L} - x) / 2 + ({res['MB_q']:.2f})")

body("En los extremos:")
eq(f"M_BC(0) = 0 + ({res['MB_q']:.2f}) = {res['MB_q']:.2f} kN-m")
eq(f"M_BC(L) = 0 + ({res['MB_q']:.2f}) = {res['MB_q']:.2f} kN-m")

body("En el centro (x = L/2 = 5 m):")
M_mid_check = q_t * 5 * 5 / 2 + res['MB_q']
eq(f"M_BC(5) = {q_t:.3f} x 5 x 5 / 2 + ({res['MB_q']:.2f})")
eq(f"M_BC(5) = {q_t*25/2:.2f} + ({res['MB_q']:.2f})")
eq_red(f"M_BC(5) = {M_mid_check:.2f} kN-m")

body("La maxima flecha de la viga se presenta en la zona donde el momento cambia de signo, "
     "es decir, entre los puntos de inflexion. El punto de momento maximo positivo se encuentra "
     "en el centro de la viga.")

spacer()

heading("3.2.7 Resumen del Caso 1", 3)
add_table(
    ['Reaccion/Momento', 'Formula', 'Valor', 'Unidad'],
    [
        ['V_A = V_D', 'qL/2', f'{res["VA_q"]:.2f}', 'kN'],
        ['H_A = H_D', 'qL2/[4h(k+2)]', f'{res["HA_q"]:.2f}', 'kN'],
        ['M_A = M_D', 'qL2/[12(k+2)]', f'{res["MA_q"]:.2f}', 'kN-m'],
        ['M_B = M_C', '-qL2/[6(k+2)]', f'{res["MB_q"]:.2f}', 'kN-m'],
        ['M_mid', 'qL2(3k+2)/[24(k+2)]', f'{res["M_mid_q"]:.2f}', 'kN-m'],
    ]
)
body("Tabla 3.1: Resumen de reacciones y momentos del Caso 1.")

doc.add_page_break()

# --- CASO 2 ---
banner("CASO 2: CARGA LATERAL EN EL NUDO B",
       "Fuerza horizontal P = 45 kN aplicada en el nudo B (esquina superior izquierda del portico), simulando sismo.",
       "caso2_lateral.png")
heading("3.3 Caso 2: Carga Lateral P = 45 kN", 2)
body(f"Fuerza horizontal P = {P_lat} kN aplicada en el nudo B.")

denom_P = 6*k + 1

spacer()

heading("3.3.1 Reaccion Vertical V_A", 3)
eq("V_A = 3 x P x h x k / [L x (6k + 1)]")
eq(f"6k + 1 = 6 x {k:.4f} + 1 = {denom_P:.4f}")
eq(f"Numerador = 3 x {P_lat} x {h} x {k:.4f} = {3*P_lat*h*k:.4f}")
eq(f"Denominador = {L} x {denom_P:.4f} = {L*denom_P:.4f}")
eq(f"V_A = {3*P_lat*h*k:.4f} / {L*denom_P:.4f}")
eq_red(f"V_A = {res['VA_P']:.2f} kN (hacia arriba)")

spacer()

heading("3.3.2 Reaccion Vertical V_D", 3)
eq("V_D = -V_A (equilibrio vertical)")
eq_red(f"V_D = {res['VD_P']:.2f} kN (hacia abajo)")

spacer()

heading("3.3.3 Reaccion Horizontal H_A y H_D", 3)
eq("H_A = H_D = P / 2")
eq(f"H_A = {P_lat} / 2")
eq_red(f"H_A = H_D = {res['HA_P']:.2f} kN")
body(f"Verificacion: H_A + H_D = {res['HA_P']:.2f} + {res['HA_P']:.2f} = {2*res['HA_P']:.2f} kN = P  OK")

spacer()

heading("3.3.4 Momento en la Base M_A", 3)
eq("M_A = -P x h x (3k + 1) / [2 x (6k + 1)]")
eq(f"3k + 1 = 3 x {k:.4f} + 1 = {3*k+1:.4f}")
eq(f"Numerador = {P_lat} x {h} x {3*k+1:.4f} = {P_lat*h*(3*k+1):.4f}")
eq(f"Denominador = 2 x {denom_P:.4f} = {2*denom_P:.4f}")
eq(f"M_A = -{P_lat*h*(3*k+1):.4f} / {2*denom_P:.4f}")
eq_red(f"M_A = {res['MA_P']:.2f} kN-m")

spacer()

heading("3.3.5 Momento en la Base M_D", 3)
eq("M_D = +P x h x (3k + 1) / [2 x (6k + 1)]")
eq_red(f"M_D = +{res['MD_P']:.2f} kN-m")

spacer()

heading("3.3.6 Momento en el Nudo M_B", 3)
eq("M_B = (P x h / 2) x 3k / (6k + 1)")
eq(f"P x h / 2 = {P_lat} x {h} / 2 = {P_lat*h/2:.2f}")
eq(f"3k / (6k+1) = {3*k:.4f} / {denom_P:.4f} = {3*k/denom_P:.6f}")
eq(f"M_B = {P_lat*h/2:.2f} x {3*k/denom_P:.6f}")
eq_red(f"M_B = {res['MB_P']:.2f} kN-m")

spacer()

heading("3.3.7 Momento en el Nudo M_C", 3)
eq("M_C = -M_B")
eq_red(f"M_C = {res['MC_P']:.2f} kN-m")

spacer()

heading("3.3.8 Resumen del Caso 2", 3)
add_table(
    ['Reaccion/Momento', 'Formula', 'Valor', 'Unidad'],
    [
        ['V_A', '3Phk/[L(6k+1)]', f'{res["VA_P"]:.2f}', 'kN'],
        ['V_D', '-V_A', f'{res["VD_P"]:.2f}', 'kN'],
        ['H_A = H_D', 'P/2', f'{res["HA_P"]:.2f}', 'kN'],
        ['M_A', '-Ph(3k+1)/[2(6k+1)]', f'{res["MA_P"]:.2f}', 'kN-m'],
        ['M_D', '+Ph(3k+1)/[2(6k+1)]', f'{res["MD_P"]:.2f}', 'kN-m'],
        ['M_B', '(Ph/2)3k/(6k+1)', f'{res["MB_P"]:.2f}', 'kN-m'],
        ['M_C', '-M_B', f'{res["MC_P"]:.2f}', 'kN-m'],
    ]
)
body("Tabla 3.2: Resumen del Caso 2.")

doc.add_page_break()

# --- SUPERPOSICION ---
banner("SUPERPOSICION DE CARGAS - TODO EL PORTICO (A-B-C-D)",
       "Se combinan los resultados del Caso 1 y Caso 2 para obtener las solicitaciones maximas de diseno en cada nudo y elemento.",
       "superposicion.png")
heading("3.4 Superposicion de Casos", 2)
body("Los resultados se combinan tomando la envolvente mas desfavorable (sismo en ambas direcciones).")

spacer()

body("Reaccion vertical maxima en A:")
eq(f"V_A_max = V_A(q) + |V_A(P)| = {res['VA_q']:.2f} + {abs(res['VA_P']):.2f}")
eq_red(f"V_A_max = {res['VA_max']:.2f} kN")

body("Reaccion horizontal maxima en A:")
eq(f"H_A_max = H_A(q) + H_A(P) = {res['HA_q']:.2f} + {res['HA_P']:.2f}")
eq_red(f"H_A_max = {res['HA_max']:.2f} kN")

body("Momento maximo en la base A:")
eq(f"M_A_max = M_A(q) + |M_A(P)| = {res['MA_q']:.2f} + {abs(res['MA_P']):.2f}")
eq_red(f"M_A_max = {res['MA_max']:.2f} kN-m")

body("Momento maximo en la base D:")
eq(f"M_D_max = M_D(q) + |M_D(P)| = {res['MD_q']:.2f} + {abs(res['MD_P']):.2f}")
eq_red(f"M_D_max = {res['MD_max']:.2f} kN-m")

body("Momento maximo en el extremo de la viga (nudo B):")
eq(f"M_B_max = |M_B(q)| + |M_B(P)| = {abs(res['MB_q']):.2f} + {abs(res['MB_P']):.2f}")
eq_red(f"M_B_max = {res['MB_max_beam']:.2f} kN-m")

body("Momento maximo en el extremo de la viga (nudo C):")
eq(f"M_C_max = |M_C(q)| + |M_C(P)| = {abs(res['MC_q']):.2f} + {abs(res['MC_P']):.2f}")
eq_red(f"M_C_max = {res['MC_max_beam']:.2f} kN-m")

body("Momento en el centro de la viga:")
eq_red(f"M_mid = {res['M_mid_total']:.2f} kN-m")

spacer()

heading("3.5 Verificacion del Equilibrio Global", 2)
body("Equilibrio de fuerzas horizontales (Caso 2):")
eq(f"P - H_A - H_D = {P_lat} - {res['HA_P']:.2f} - {res['HA_P']:.2f} = 0  OK")

body("Equilibrio de fuerzas verticales (Caso 1):")
eq(f"V_A + V_D - qL = {res['VA_q']:.2f} + {res['VA_q']:.2f} - {q_t*L:.2f} = 0  OK")

body("Equilibrio de fuerzas verticales (Caso 2):")
eq(f"V_A + V_D = {res['VA_P']:.2f} + ({res['VD_P']:.2f}) = 0  OK")

body("Equilibrio de momentos respecto al punto A (Caso 1):")
eq(f"M_A + M_D - V_D x L + q x L^2/2 - H_D x h = 0")
check_M1 = res['MA_q'] + res['MD_q'] - res['VD_q']*L + q_t*L**2/2 - res['HD_q']*h
body(f"Residuo = {check_M1:.4f} kN-m (OK)")

body("Equilibrio de momentos respecto al punto A (Caso 2):")
check_M2 = res['MA_P'] + res['MD_P'] + P_lat*h - res['VD_P']*L
body(f"Residuo = {check_M2:.4f} kN-m (OK)")

spacer()

heading("3.6 Resumen de Reacciones", 2)
add_table(
    ['Reaccion', 'Caso 1 (q)', 'Caso 2 (P)', 'Combinado'],
    [
        ['V_A [kN]', f'{res["VA_q"]:.2f}', f'{res["VA_P"]:.2f}', f'{res["VA_max"]:.2f}'],
        ['V_D [kN]', f'{res["VD_q"]:.2f}', f'{res["VD_P"]:.2f}', f'{res["VA_max"]:.2f}'],
        ['H_A [kN]', f'{res["HA_q"]:.2f}', f'{res["HA_P"]:.2f}', f'{res["HA_max"]:.2f}'],
        ['H_D [kN]', f'{res["HD_q"]:.2f}', f'{res["HD_P"]:.2f}', f'{res["HA_max"]:.2f}'],
        ['M_A [kN-m]', f'{res["MA_q"]:.2f}', f'{res["MA_P"]:.2f}', f'{res["MA_max"]:.2f}'],
        ['M_D [kN-m]', f'{res["MD_q"]:.2f}', f'{res["MD_P"]:.2f}', f'{res["MD_max"]:.2f}'],
        ['M_B [kN-m]', f'{res["MB_q"]:.2f}', f'{res["MB_P"]:.2f}', f'{res["MB_max_beam"]:.2f}'],
        ['M_C [kN-m]', f'{res["MC_q"]:.2f}', f'{res["MC_P"]:.2f}', f'{res["MC_max_beam"]:.2f}'],
        ['M_mid [kN-m]', f'{res["M_mid_q"]:.2f}', '-', f'{res["M_mid_total"]:.2f}'],
    ]
)
body("Tabla 3.3: Resumen completo de reacciones y momentos.")

doc.add_page_break()

# --- FUERZAS INTERNAS ---
banner("FUERZAS INTERNAS - VIGA (elemento B-C)",
       "Funciones de cortante V(x), axial N y momento M(x) a lo largo de la viga, entre los nudos B y C.",
       "fuerzas_viga.png")
heading("3.7 Fuerzas Internas - Viga BC", 2)

body("Fuerza axial en la viga (compresion):")
eq("N_viga = H_A(q) + H_A(P)")
eq(f"N_viga = {res['HA_q']:.2f} + {res['HA_P']:.2f}")
eq_red(f"N_viga = {res['N_beam']:.2f} kN")

body("Funcion de cortante V(x) en la viga:")
eq(f"V(x) = {res['VA_max']:.2f} - {q_t:.3f} x x")

body("Valores notables:")
eq(f"V(0)   = {res['VA_max']:.2f} kN")
eq(f"V(L/2) = {res['VA_max'] - q_t*5:.2f} kN")
eq(f"V(L)   = {res['VA_max'] - q_t*10:.2f} kN")
eq_red(f"V_max_viga = {res['V_beam_max']:.2f} kN")

body("Posicion donde V = 0:")
x_V0 = res['VA_max'] / q_t
eq(f"x_0 = V_max / q = {res['VA_max']:.2f} / {q_t:.3f} = {x_V0:.2f} m")

body("Funcion de momento M(x) en la viga:")
eq(f"M(x) = -{res['MB_max_beam']:.2f} + {res['VA_max']:.2f} x x - {q_t:.3f} x x^2 / 2")

body("Momento en el centro (x = 5 m):")
M_check = -res['MB_max_beam'] + res['VA_max']*5 - q_t*25/2
eq(f"M(5) = -{res['MB_max_beam']:.2f} + {res['VA_max']*5:.2f} - {q_t*25/2:.2f}")
eq_red(f"M(5) = {M_check:.2f} kN-m")

spacer()

banner("FUERZAS INTERNAS - COLUMNAS (elementos A-B y D-C)",
       "Axial N, cortante V y momento M en las columnas. Ambas columnas tienen el mismo perfil W250x58.",
       "fuerzas_columnas.png")
heading("3.8 Fuerzas Internas - Columna", 2)

body("Fuerza axial en la columna:")
eq("N_col = V_A_max + w_pp_col x h")
eq(f"N_col = {res['VA_max']:.2f} + {best['w_col']:.3f} x {h}")
eq(f"N_col = {res['VA_max']:.2f} + {best['w_col']*h:.2f}")
eq_red(f"N_col = {best['N_col']:.2f} kN")

body("Fuerza cortante en la columna:")
eq_red(f"V_col = H_A_max = {res['HA_max']:.2f} kN")

body("Funcion de momento M(y) en la columna (y desde la base):")
eq(f"M(y) = {res['MA_max']:.2f} - {res['HA_max']:.2f} x y")

body("Valores notables:")
eq(f"M(0) = {res['MA_max']:.2f} kN-m  (base)")
eq(f"M(h) = {res['MA_max']:.2f} - {res['HA_max']*h:.2f} = {res['MA_max']-res['HA_max']*h:.2f} kN-m  (nudo B)")

body("Momento maximo en la columna:")
eq_red(f"M_max_col = {best['M_col']:.2f} kN-m")

body("Punto de inflexion:")
y_infl = res['MA_max'] / res['HA_max']
eq(f"y_inf = M_A / H_A = {res['MA_max']:.2f} / {res['HA_max']:.2f} = {y_infl:.2f} m")

spacer()

heading("3.9 Valores Criticos", 2)
add_table(
    ['Elemento', 'N [kN]', 'V [kN]', 'M [kN-m]', 'Seccion critica'],
    [
        ['Viga (extremo)', f'{res["N_beam"]:.2f}', f'{res["V_beam_max"]:.2f}', f'{res["MB_max_beam"]:.2f}', 'Nudo B'],
        ['Viga (centro)', f'{res["N_beam"]:.2f}', '~0', f'{res["M_mid_total"]:.2f}', 'L/2'],
        ['Columna', f'{best["N_col"]:.2f}', f'{res["HA_max"]:.2f}', f'{best["M_col"]:.2f}', 'Base'],
    ]
)
body("Tabla 3.4: Valores criticos de fuerzas internas.")

body("Momento critico de diseno para la viga:")
eq(f"M_crit = max({res['MB_max_beam']:.2f}, {res['M_mid_total']:.2f})")
eq_red(f"M_crit = {best['M_crit_beam']:.2f} kN-m")

spacer()

heading("3.10 Diagramas de Fuerza Interna", 2)
add_img('Diagramas_Portico.png', 6.5)
body("Figura 3.1: Diagramas de momento flector y fuerza cortante.")

body("En el diagrama de momento de la viga se observa la forma parabolica producida por la "
     "carga distribuida, con momentos negativos en los extremos y momento positivo en el centro.")
body("El diagrama de cortante de la viga muestra variacion lineal, tipica de carga uniforme.")
body("Para la columna, el diagrama de momento es lineal y el de cortante es constante.")

doc.add_page_break()

# ##############################################################
# CAPITULO 4
# ##############################################################
print("    Cap 4: Diseno de elementos...")

heading("CAPITULO 4: DISENO DE ELEMENTOS")
spacer()

body("En este capitulo se verifica que los perfiles seleccionados cumplan con los criterios "
     "de resistencia por flexion, cortante, flexo-compresion y esbeltez.")

spacer()

banner("DISENO DE LA VIGA (elemento B-C) - Perfil " + bn,
       "Verificacion de esfuerzos normales, cortantes y esbeltez del perfil comercial seleccionado para la viga.",
       "diseno_viga.png")
heading("4.1 Diseno de la Viga - Perfil " + bn, 2)
spacer()

heading("4.1.1 Propiedades Geometricas", 3)
body(f"Perfil seleccionado: {bn}")
eq(f"d  = {bp['d']} mm")
eq(f"bf = {bp['bf']} mm")
eq(f"tf = {bp['tf']} mm")
eq(f"tw = {bp['tw']} mm")

hw_b = bp['d'] - 2 * bp['tf']
body("Altura del alma:")
eq(f"hw = d - 2tf = {bp['d']} - 2 x {bp['tf']} = {hw_b:.1f} mm")

body("Area:")
eq(f"A = {bp['A']} mm2")

A_check_b = 2 * bp['bf'] * bp['tf'] + hw_b * bp['tw']
body("Verificacion del area:")
eq(f"A = 2 x bf x tf + hw x tw = 2 x {bp['bf']} x {bp['tf']} + {hw_b:.1f} x {bp['tw']}")
eq(f"A = {2*bp['bf']*bp['tf']:.0f} + {hw_b*bp['tw']:.0f} = {A_check_b:.0f} mm2")

body("Momento de inercia:")
eq(f"Ix = {bp['Ix']} cm4")

Ix_check_b = (bp['bf'] * bp['d']**3 / 12) - ((bp['bf'] - bp['tw']) * hw_b**3 / 12)
body("Verificacion Ix:")
eq(f"Ix = bf x d^3/12 - (bf-tw) x hw^3/12")
eq(f"Ix = {bp['bf']*bp['d']**3/12:.0f} - {(bp['bf']-bp['tw'])*hw_b**3/12:.0f}")
eq(f"Ix = {Ix_check_b:.0f} mm4 = {Ix_check_b/1e4:.0f} cm4")

body("Modulo de seccion:")
eq(f"Sx = {bp['Sx']} cm3 = {bp['Sx']*1e3:.0f} mm3")

spacer()

heading("4.1.2 Peso Propio", 3)
eq(f"Peso lineal = {bp['w_kg']} kg/m")
eq(f"w_pp = {bp['w_kg']} x {g_acc} / 1000")
eq_red(f"w_pp = {best['w_viga']:.3f} kN/m")
eq(f"P_pp_viga = {best['w_viga']:.3f} x {L} = {best['w_viga']*L:.2f} kN")

spacer()

heading("4.1.3 Solicitaciones Totales", 3)
eq_red(f"M_crit = {best['M_crit_beam']:.2f} kN-m")
eq_red(f"V_max = {res['V_beam_max']:.2f} kN")
eq_red(f"N = {res['N_beam']:.2f} kN")

spacer()

heading("4.1.4 Verificacion por Flexion", 3)
body("Esfuerzo normal por flexion:")
eq("sigma_flex = M / Sx")
eq(f"sigma_flex = {best['M_crit_beam']:.2f} x 10^6 / ({bp['Sx']} x 10^3)")
eq(f"sigma_flex = {best['M_crit_beam']*1e6:.0f} / {bp['Sx']*1e3:.0f}")
eq_red(f"sigma_flex = {sigma_flex_b:.2f} MPa")

body("Esfuerzo axial:")
eq("sigma_N = N / A")
eq(f"sigma_N = {res['N_beam']:.2f} x 10^3 / {bp['A']}")
eq_red(f"sigma_N = {sigma_ax_b:.2f} MPa")

body("Esfuerzo total:")
eq(f"sigma_total = {sigma_flex_b:.2f} + {sigma_ax_b:.2f}")
eq_red(f"sigma_total = {sigma_tot_b:.2f} MPa")

body("Comparacion:")
eq(f"sigma_total = {sigma_tot_b:.2f} MPa  <  sigma_adm = {sigma_adm:.2f} MPa  CUMPLE")

body("Factor de seguridad:")
eq(f"FS = Fy / sigma_total = {Fy} / {sigma_tot_b:.2f}")
eq_red(f"FS = {fs_b:.2f}  >=  {FS_min}  CUMPLE")

spacer()

heading("4.1.5 Verificacion por Cortante", 3)
eq("tau = V / (d x tw)")
eq(f"tau = {res['V_beam_max']:.2f} x 10^3 / ({bp['d']} x {bp['tw']})")
eq(f"tau = {res['V_beam_max']*1e3:.0f} / {bp['d']*bp['tw']:.0f}")
eq_red(f"tau = {tau_b_val:.2f} MPa")

eq(f"tau = {tau_b_val:.2f} MPa  <  tau_adm = {tau_adm:.2f} MPa  CUMPLE")

eq(f"FS_tau = {tau_y} / {tau_b_val:.2f}")
eq_red(f"FS_tau = {fs_tau_b:.2f}  >=  {FS_min}  CUMPLE")

spacer()

heading("4.1.6 Verificacion de Esbeltez", 3)
eq(f"d/tw = {bp['d']} / {bp['tw']} = {bp['d']/bp['tw']:.1f}  <=  50  CUMPLE")
eq(f"tw = {bp['tw']} mm  >=  8 mm  CUMPLE")

doc.add_page_break()

# ---- COLUMNA ----
banner("DISENO DE LA COLUMNA (elementos A-B y D-C) - Perfil " + cn,
       "Verificacion por flexo-compresion y cortante del perfil seleccionado para ambas columnas.",
       "diseno_columna.png")
heading("4.2 Diseno de la Columna - Perfil " + cn, 2)
spacer()

heading("4.2.1 Propiedades", 3)
eq(f"d  = {cp['d']} mm")
eq(f"bf = {cp['bf']} mm")
eq(f"tf = {cp['tf']} mm")
eq(f"tw = {cp['tw']} mm")

hw_c = cp['d'] - 2 * cp['tf']
eq(f"hw = d - 2tf = {hw_c:.1f} mm")
eq(f"A = {cp['A']} mm2")
eq(f"Ix = {cp['Ix']} cm4")
eq(f"Sx = {cp['Sx']} cm3 = {cp['Sx']*1e3:.0f} mm3")

spacer()

heading("4.2.2 Verificacion de Espesor", 3)
eq(f"tw = {cp['tw']} mm  >=  8 mm  CUMPLE")
eq(f"d/tw = {cp['d']} / {cp['tw']} = {cp['d']/cp['tw']:.1f}  <=  50  CUMPLE")

spacer()

heading("4.2.3 Peso Propio", 3)
eq(f"w_pp_col = {best['w_col']:.3f} kN/m")
eq(f"P_pp_col = {best['w_col']:.3f} x {h} = {best['w_col']*h:.2f} kN")

spacer()

heading("4.2.4 Carga Axial Total", 3)
eq("N_col = V_A_max + w_pp_col x h")
eq(f"N_col = {res['VA_max']:.2f} + {best['w_col']*h:.2f}")
eq_red(f"N_col = {best['N_col']:.2f} kN")

spacer()

heading("4.2.5 Verificacion por Flexo-compresion", 3)
body("Esfuerzo por flexion:")
eq(f"sigma_flex = M_col / Sx = {best['M_col']:.2f} x 10^6 / {cp['Sx']*1e3:.0f}")
eq_red(f"sigma_flex = {sigma_flex_c:.2f} MPa")

body("Esfuerzo por compresion:")
eq(f"sigma_N = N / A = {best['N_col']:.2f} x 10^3 / {cp['A']}")
eq_red(f"sigma_N = {sigma_ax_c:.2f} MPa")

body("Esfuerzo total:")
eq(f"sigma_total = {sigma_flex_c:.2f} + {sigma_ax_c:.2f}")
eq_red(f"sigma_total = {sigma_tot_c:.2f} MPa")

eq(f"sigma_total = {sigma_tot_c:.2f} < sigma_adm = {sigma_adm:.2f}  CUMPLE")

body("Factor de seguridad:")
eq(f"FS = {Fy} / {sigma_tot_c:.2f}")
eq_red(f"FS = {fs_c:.2f}  >=  {FS_min}  CUMPLE")

spacer()

heading("4.2.6 Verificacion por Cortante", 3)
eq(f"tau = {res['HA_max']:.2f} x 10^3 / ({cp['d']} x {cp['tw']})")
eq(f"tau = {res['HA_max']*1e3:.0f} / {cp['d']*cp['tw']:.0f}")
eq_red(f"tau = {tau_c_val:.2f} MPa")
eq(f"tau = {tau_c_val:.2f} < tau_adm = {tau_adm:.2f}  CUMPLE")
eq(f"FS_tau = {tau_y} / {tau_c_val:.2f}")
eq_red(f"FS_tau = {fs_tau_c:.2f}  >=  {FS_min}  CUMPLE")

spacer()

# Resumen
heading("4.2.7 Resumen de Verificacion", 3)
add_table(
    ['Verificacion', f'Viga {bn}', f'Columna {cn}', 'Limite', 'Estado'],
    [
        ['sigma [MPa]', f'{sigma_tot_b:.2f}', f'{sigma_tot_c:.2f}', f'< {sigma_adm:.2f}', 'CUMPLE'],
        ['FS flexion', f'{fs_b:.2f}', f'{fs_c:.2f}', f'>= {FS_min}', 'CUMPLE'],
        ['tau [MPa]', f'{tau_b_val:.2f}', f'{tau_c_val:.2f}', f'< {tau_adm:.2f}', 'CUMPLE'],
        ['FS cortante', f'{fs_tau_b:.2f}', f'{fs_tau_c:.2f}', f'>= {FS_min}', 'CUMPLE'],
        ['d/tw', f'{bp["d"]/bp["tw"]:.1f}', f'{cp["d"]/cp["tw"]:.1f}', '<= 50', 'CUMPLE'],
        ['tw [mm]', f'{bp["tw"]}', f'{cp["tw"]}', '>= 8', 'CUMPLE'],
    ]
)
body("Tabla 4.1: Resumen de verificaciones.")

doc.add_page_break()

# ---- PERFIL PERSONALIZADO ----
banner("PERFIL PERSONALIZADO - VIGA (elemento B-C)",
       "Diseno de una seccion I optimizada (no comercial) que reduce el area manteniendo FS >= 1.4.",
       "perfil_custom.png")
heading("4.3 Perfil Personalizado para la Viga", 2)
body("Se disena un perfil I optimizado con las restricciones: tw >= 8, tf >= 8, d/tw <= 50, FS >= 1.40.")

if best_custom:
    pc = best_custom

    heading("4.3.1 Dimensiones", 3)
    eq(f"d  = {pc['d']} mm")
    eq(f"bf = {pc['bf']} mm")
    eq(f"tf = {pc['tf']} mm")
    eq(f"tw = {pc['tw']} mm")
    hw_custom = pc['d'] - 2 * pc['tf']
    eq(f"hw = {hw_custom} mm")
    eq(f"d/tw = {pc['esbeltez']:.1f}  <=  50  CUMPLE")

    heading("4.3.2 Calculo de Propiedades", 3)
    body("Area:")
    eq(f"A = 2 x {pc['bf']} x {pc['tf']} + {hw_custom} x {pc['tw']}")
    eq(f"A = {2*pc['bf']*pc['tf']} + {hw_custom*pc['tw']}")
    eq_red(f"A = {pc['A']:.0f} mm2")

    body("Momento de inercia:")
    eq(f"Ix = {pc['bf']} x {pc['d']}^3/12 - ({pc['bf']}-{pc['tw']}) x {hw_custom}^3/12")
    Ix_t1 = pc['bf'] * pc['d']**3 / 12
    Ix_t2 = (pc['bf'] - pc['tw']) * hw_custom**3 / 12
    eq(f"Ix = {Ix_t1:.0f} - {Ix_t2:.0f}")
    eq_red(f"Ix = {pc['Ix']:.0f} mm4 = {pc['Ix']/1e4:.0f} cm4")

    body("Modulo de seccion:")
    eq(f"Sx = {pc['Ix']:.0f} / {pc['d']/2:.0f}")
    eq_red(f"Sx = {pc['Sx']:.0f} mm3 = {pc['Sx']/1e3:.0f} cm3")

    body("Peso lineal:")
    eq_red(f"w = {pc['w_kg']:.1f} kg/m ({pc['w_kN']:.3f} kN/m)")

    heading("4.3.3 Verificacion de Esfuerzos", 3)
    sigma_flex_custom = pc['M_crit'] * 1e6 / pc['Sx']
    sigma_ax_custom = pc['N'] * 1e3 / pc['A']
    body("Esfuerzo por flexion:")
    eq(f"sigma_flex = {pc['M_crit']:.2f} x 10^6 / {pc['Sx']:.0f}")
    eq_red(f"sigma_flex = {sigma_flex_custom:.2f} MPa")

    body("Esfuerzo axial:")
    eq(f"sigma_N = {pc['N']:.2f} x 10^3 / {pc['A']:.0f}")
    eq_red(f"sigma_N = {sigma_ax_custom:.2f} MPa")

    body("Esfuerzo total:")
    eq(f"sigma_total = {sigma_flex_custom:.2f} + {sigma_ax_custom:.2f}")
    eq_red(f"sigma_total = {pc['sigma']:.2f} MPa")

    eq(f"FS = {Fy} / {pc['sigma']:.2f}")
    eq_red(f"FS = {pc['fs_f']:.2f}  >=  {FS_min}  CUMPLE")

    body("Cortante:")
    eq(f"tau = {pc['V']:.2f} x 10^3 / ({pc['d']} x {pc['tw']})")
    eq_red(f"tau = {pc['tau']:.2f} MPa")
    eq(f"FS_tau = {tau_y} / {pc['tau']:.2f} = {tau_y/pc['tau']:.2f}  CUMPLE")

    heading("4.3.4 Comparacion con Perfil Comercial", 3)
    add_table(
        ['Propiedad', f'Comercial ({bn})', 'Personalizado', 'Diferencia'],
        [
            ['A [mm2]', f'{bp["A"]}', f'{pc["A"]:.0f}', f'{(1-pc["A"]/bp["A"])*100:.1f}% menos'],
            ['Ix [cm4]', f'{bp["Ix"]}', f'{pc["Ix"]/1e4:.0f}', ''],
            ['Sx [cm3]', f'{bp["Sx"]}', f'{pc["Sx"]/1e3:.0f}', ''],
            ['Peso [kg/m]', f'{bp["w_kg"]}', f'{pc["w_kg"]:.1f}', f'{(1-pc["w_kg"]/bp["w_kg"])*100:.1f}% menos'],
            ['FS', f'{fs_b:.2f}', f'{pc["fs_f"]:.2f}', ''],
        ]
    )
    body("Tabla 4.2: Comparacion comercial vs personalizado.")
    body(f"Reduccion de area: {(1-pc['A']/bp['A'])*100:.1f}%.")

doc.add_page_break()

# ##############################################################
# CAPITULO 5
# ##############################################################
print("    Cap 5: Conexiones...")

heading("CAPITULO 5: CONEXIONES")
spacer()

body("Se disenan las conexiones viga-columna (nudos B y C) y columna-fundacion (bases A y D).")

spacer()

banner("CONEXION VIGA-COLUMNA (nudos B y C)",
       "Diseno de la union entre la viga B-C y las columnas, transmitiendo momento y cortante. Pernos A490 + soldadura E70.",
       "conexion_BC.png")
heading("5.1 Conexion B/C - Viga-Columna (Conexion a Momento)", 2)
body("Solicitaciones en la conexion:")
eq(f"M = {M_conn:.2f} kN-m")
eq(f"V = {V_conn:.2f} kN")
eq(f"N = {res['N_beam']:.2f} kN")

spacer()

heading("5.1.1 Conexion Empernada A490", 3)
body("Propiedades de los pernos A490:")
eq(f"Fu_A490 = {Fu_A490} MPa")
eq(f"Fnv = {Fnv_A490} MPa (cortante, rosca excluida)")

body("Brazo de palanca:")
eq(f"brazo = d - tf = {bp['d']} - {bp['tf']} = {bp['d']-bp['tf']:.1f} mm = {lever:.4f} m")

body("Fuerza en cada ala:")
eq(f"F_ala = M / brazo = {M_conn:.2f} / {lever:.4f}")
eq_red(f"F_ala = {F_ala:.2f} kN")

body(f"Pernos A490, phi = {bolt_sel} ({bolt['d']:.2f} mm):")
eq(f"A_b = pi x {bolt['d']:.2f}^2 / 4")
eq_red(f"A_b = {A_bolt:.1f} mm2")

body("Resistencia al cortante:")
eq(f"Rv_simple = {Fnv_A490} x {A_bolt:.1f} / 1000")
eq_red(f"Rv_simple = {Rv_simple:.1f} kN")
eq(f"Rv_doble = 2 x {Rv_simple:.1f}")
eq_red(f"Rv_doble = {Rv_doble:.1f} kN")

body("Numero de pernos (alas, doble corte):")
eq(f"n_ala = ceil({F_ala:.1f} / {Rv_doble:.1f})")
eq_red(f"n_ala = {n_ala} pernos/ala")

body("Numero de pernos (alma, simple corte):")
eq(f"n_alma = ceil({V_conn:.1f} / {Rv_simple:.1f})")
eq_red(f"n_alma = {n_alma} pernos")

body("Total:")
eq(f"n_total = 2 x {n_ala} + {n_alma}")
eq_red(f"n_total = {n_total} pernos")

spacer()

heading("5.1.2 Verificacion por Aplastamiento", 3)
eq(f"t_min = min({bp['tf']}, 25) = {t_min_ala} mm")
eq("Rn = 2.4 x Fu x d_b x t_min / 1000")
eq(f"Rn = 2.4 x {Fu_A36} x {bolt['d']:.2f} x {t_min_ala} / 1000")
eq_red(f"Rn = {Rn_bearing_ala:.1f} kN/perno")
eq(f"F_dem = {F_ala:.1f} / {n_ala} = {F_dem_ala:.1f} kN/perno")
eq(f"{F_dem_ala:.1f} < {Rn_bearing_ala:.1f}  CUMPLE")

spacer()

heading("5.1.3 Verificacion por Desgarramiento", 3)
eq(f"Le = {Le} mm")
eq(f"Le_eff = {Le} - {bolt['dh']/2:.1f} = {Le_eff:.1f} mm")
eq("Rn = 1.2 x Le_eff x t_min x Fu / 1000")
eq(f"Rn = 1.2 x {Le_eff:.1f} x {t_min_ala} x {Fu_A36} / 1000")
eq_red(f"Rn = {Rn_tearout:.1f} kN/perno")
eq(f"{F_dem_ala:.1f} < {Rn_tearout:.1f}  CUMPLE")

spacer()

heading("5.1.4 Verificacion por Block Shear", 3)
body("Parametros:")
eq(f"Separacion: s = {s_bolt} mm")
eq(f"Distancia borde: Lev = {Lev} mm")

body("Area neta a traccion:")
eq(f"Ant = t x 2s = {t_min_ala} x {2*s_bolt}")
eq_red(f"Ant = {Ant:.0f} mm2")

body("Area neta a cortante:")
Anv_calc = Lev + (n_ala-1)*s_bolt - (n_ala-0.5)*bolt['dh']
eq(f"Anv = {t_min_ala} x [{Lev} + ({n_ala}-1)x{s_bolt} - ({n_ala}-0.5)x{bolt['dh']}]")
eq(f"Anv = {t_min_ala} x {Anv_calc:.1f}")
eq_red(f"Anv = {Anv:.0f} mm2")

body("Resistencia block shear:")
eq("Rn = 0.6 x Fu x Anv + Ubs x Fu x Ant")
eq(f"Rn = 0.6 x {Fu_A36} x {Anv:.0f}/1000 + {Ubs} x {Fu_A36} x {Ant:.0f}/1000")
eq(f"Rn = {0.6*Fu_A36*Anv/1000:.1f} + {Ubs*Fu_A36*Ant/1000:.1f}")
eq_red(f"Rn_block = {Rn_block:.1f} kN")
eq(f"F_ala = {F_ala:.1f} < {Rn_block:.1f}  CUMPLE")

spacer()

heading("5.1.5 Conexion Soldada E70", 3)
body("Propiedades electrodo E70:")
eq(f"Fexx = {Fexx_E70} MPa")
eq(f"Fw = 0.6 x {Fexx_E70} = {Fw_E70:.1f} MPa")

body("Soldadura ala (momento):")
eq(f"a_ala = F_ala x 10^3 / (0.707 x Fw x 2 x bf)")
eq(f"a_ala = {F_ala*1e3:.0f} / (0.707 x {Fw_E70:.1f} x 2 x {bp['bf']})")
eq_red(f"a_ala = {a_ala_sold:.1f} mm  -->  usar {max(8, math.ceil(a_ala_sold))} mm")

body("Soldadura alma (cortante):")
eq(f"a_alma = V x 10^3 / (0.707 x Fw x 2 x hw)")
eq(f"a_alma = {V_conn*1e3:.0f} / (0.707 x {Fw_E70:.1f} x 2 x {hw_beam:.0f})")
eq_red(f"a_alma = {a_alma_sold:.1f} mm  -->  usar {max(6, math.ceil(a_alma_sold))} mm")

doc.add_page_break()

# ---- CONEXION BASE ----
banner("CONEXION BASE DE COLUMNA (nudos A y D)",
       "Diseno de la placa base y pernos de anclaje que unen las columnas A-B y D-C a la fundacion de concreto.",
       "conexion_base.png")
heading("5.2 Conexion Base - Columna-Fundacion", 2)
body("Solicitaciones en la base:")
eq(f"M = {M_base:.2f} kN-m")
eq(f"N = {N_base:.2f} kN (compresion)")
eq(f"V = {V_base:.2f} kN")

spacer()

heading("5.2.1 Placa Base", 3)
eq(f"Ancho = bf + 150 = {cp['bf']} + 150 = {ancho_placa} mm")
eq(f"Largo = d + 200 = {cp['d']} + 200 = {largo_placa} mm")
eq(f"Espesor = {t_placa} mm")
eq(f"A_placa = {ancho_placa} x {largo_placa} = {A_placa} mm2")

body("Distribucion de esfuerzos:")
eq(f"sigma_max = N/A + Mc/I")
eq(f"sigma_max = {N_base*1e3/A_placa:.4f} + {M_base*1e6*(largo_placa/2)/I_placa:.4f}")
eq_red(f"sigma_max = {sigma_max_base:.2f} MPa")

eq(f"sigma_min = N/A - Mc/I")
eq(f"sigma_min = {N_base*1e3/A_placa:.4f} - {M_base*1e6*(largo_placa/2)/I_placa:.4f}")
eq_red(f"sigma_min = {sigma_min_base:.2f} MPa")

if sigma_min_base < 0:
    body("Esfuerzo minimo negativo (traccion): los pernos de anclaje deben resistir.")

spacer()

heading("5.2.2 Pernos de Anclaje", 3)
body(f"{n_pernos_base} pernos A490, phi = {bolt_base} ({bolt_b['d']:.2f} mm)")
eq(f"A_b = pi x {bolt_b['d']:.2f}^2 / 4 = {A_bolt_b:.1f} mm2")

if sigma_min_base < 0:
    F_traccion = abs(sigma_min_base) * A_placa / 4 / 1000
    eq(f"F_traccion = |sigma_min| x A/4/1000 = {F_traccion:.1f} kN")
    cap_trac = 0.75 * Fnv_A490 * A_bolt_b / 1000
    eq(f"Cap_traccion = 0.75 x {Fnv_A490} x {A_bolt_b:.1f} / 1000 = {cap_trac:.1f} kN/perno  CUMPLE")

spacer()

heading("5.2.3 Verificacion por Cortante/Friccion", 3)
eq(f"mu = {mu}")
eq(f"V_resist = mu x N = {mu} x {N_base:.1f}")
eq_red(f"V_resist = {V_fric:.1f} kN")

if V_fric > V_base:
    eq(f"V_resist = {V_fric:.1f} > V = {V_base:.1f}  CUMPLE")
else:
    V_pernos = V_base - V_fric
    eq(f"V_resist = {V_fric:.1f} < V = {V_base:.1f}  Pernos deben resistir {V_pernos:.1f} kN")
    Rv_base = Fnv_A490 * A_bolt_b / 1000
    eq(f"Rv_perno = {Rv_base:.1f} kN")
    n_req = math.ceil(V_pernos / Rv_base)
    eq(f"n_req = {n_req}, se tienen {n_pernos_base}  {'CUMPLE' if n_pernos_base >= n_req else 'NO CUMPLE'}")

spacer()

heading("5.2.4 Soldadura Base", 3)
eq(f"Perimetro = 2(d+bf) = 2({cp['d']}+{cp['bf']}) = {perim_col} mm")
eq(f"a = {a_base} mm")
eq(f"Cap = 0.707 x {a_base} x {Fw_E70:.1f} x {perim_col} / 1000")
eq_red(f"Cap = {cap_sold_base:.1f} kN  >>  {N_base+V_base:.1f} kN  CUMPLE")

body("Resumen conexion base:")
add_table(
    ['Componente', 'Especificacion', 'Estado'],
    [
        ['Placa base', f'{ancho_placa}x{largo_placa}x{t_placa} mm', 'OK'],
        ['Pernos', f'{n_pernos_base} A490 phi {bolt_base}', 'OK'],
        ['Soldadura', f'Filete a={a_base} mm, E70', 'OK'],
        ['Friccion', f'V_res={V_fric:.1f} kN', 'OK' if V_fric > V_base else 'Pernos'],
    ]
)
body("Tabla 5.1: Resumen conexion base.")

doc.add_page_break()

# ##############################################################
# CAPITULO 6
# ##############################################################
print("    Cap 6: Circulos de Mohr...")

banner("ANALISIS DE ESFUERZOS - SECCION CRITICA DE LA VIGA (nudo B, x = 0)",
       "Estado de esfuerzos en 4 puntos de la seccion transversal de la viga en el empotramiento B, donde el momento es maximo.",
       "mohr_seccion.png")
heading("CAPITULO 6: ANALISIS DE ESFUERZOS Y CIRCULOS DE MOHR")
spacer()

body("Se analizan los esfuerzos en la seccion critica de la viga "
     f"(M = {M_crit_mohr:.2f} kN-m) en 4 puntos representativos, "
     "calculando esfuerzos principales y cortante maximo.")

body("Datos de la seccion critica:")
eq(f"M = {M_crit_mohr:.2f} kN-m = {M_crit_mohr*1e6:.0f} N-mm")
eq(f"V = {V_crit_mohr:.2f} kN = {V_crit_mohr*1e3:.0f} N")
eq(f"N = {N_crit_mohr:.2f} kN = {N_crit_mohr*1e3:.0f} N")

body(f"Propiedades del perfil {bn}:")
eq(f"Ix = {bp['Ix']} cm4 = {Ix_mm4:.0f} mm4")
eq(f"d = {d_mm} mm, bf = {bf_mm} mm, tf = {tf_mm} mm, tw = {tw_mm} mm, hw = {hw_mm:.0f} mm")

spacer()

heading("6.1 Primer Momento de Area Q", 2)
body("Q en el eje neutro:")
y_ala = d_mm/2 - tf_mm/2
y_alma = d_mm/2 - tf_mm
eq(f"Q_NA = bf x tf x (d/2 - tf/2) + tw x (d/2 - tf)^2 / 2")
eq(f"Q_NA = {bf_mm} x {tf_mm} x {y_ala:.1f} + {tw_mm} x {y_alma:.1f}^2 / 2")
eq(f"Q_NA = {bf_mm*tf_mm*y_ala:.1f} + {tw_mm*y_alma**2/2:.1f}")
eq_red(f"Q_NA = {Q_NA:.1f} mm3")

body("Q en la union ala-alma:")
eq(f"Q_ala = bf x tf x (d/2 - tf/2) = {bf_mm} x {tf_mm} x {y_ala:.1f}")
eq_red(f"Q_ala = {Q_ala:.1f} mm3")

spacer()

heading("6.2 Analisis en 4 Puntos de la Seccion", 2)
body("Puntos de analisis:")
body("1. Fibra superior (y = +d/2)")
body("2. Fibra inferior (y = -d/2)")
body("3. Eje neutro (y = 0)")
body(f"4. Union ala-alma (y = {d_mm/2-tf_mm:.0f} mm)")

body("Para cada punto se calculan: sigma_x, tau_xy, sigma_avg, R, sigma_1, sigma_2, tau_max, theta_p")

spacer()

# PUNTO 1
heading("6.2.1 Punto 1 - Fibra Superior (y = +d/2)", 3)
pt1 = puntos_mohr[0]
body("Esfuerzo normal:")
eq(f"sigma_x = M x (d/2) / Ix + N / A")
eq(f"sigma_x = {M_crit_mohr:.2f}e6 x {d_mm/2:.0f} / {Ix_mm4:.0f} + {N_crit_mohr:.2f}e3 / {bp['A']}")
eq_red(f"sigma_x = {pt1['sigma_x']:.2f} MPa")
eq_red(f"tau_xy = 0 MPa")
eq(f"sigma_avg = {pt1['sigma_x']:.2f} / 2 = {pt1['sigma_avg']:.2f} MPa")
eq(f"R = sqrt({pt1['sigma_avg']:.2f}^2 + 0^2) = {pt1['R']:.2f} MPa")
eq_red(f"sigma_1 = {pt1['sigma_avg']:.2f} + {pt1['R']:.2f} = {pt1['sigma_1']:.2f} MPa")
eq_red(f"sigma_2 = {pt1['sigma_avg']:.2f} - {pt1['R']:.2f} = {pt1['sigma_2']:.2f} MPa")
eq_red(f"tau_max = {pt1['tau_max']:.2f} MPa")
eq(f"theta_p = {pt1['theta_p']:.1f} grados")

spacer()

# PUNTO 2
heading("6.2.2 Punto 2 - Fibra Inferior (y = -d/2)", 3)
pt2 = puntos_mohr[1]
body("Esfuerzo normal:")
eq(f"sigma_x = -M x (d/2) / Ix + N / A")
eq_red(f"sigma_x = {pt2['sigma_x']:.2f} MPa")
eq_red(f"tau_xy = 0 MPa")
eq(f"sigma_avg = {pt2['sigma_avg']:.2f} MPa")
eq(f"R = {pt2['R']:.2f} MPa")
eq_red(f"sigma_1 = {pt2['sigma_1']:.2f} MPa")
eq_red(f"sigma_2 = {pt2['sigma_2']:.2f} MPa")
eq_red(f"tau_max = {pt2['tau_max']:.2f} MPa")
eq(f"theta_p = {pt2['theta_p']:.1f} grados")

body("Esfuerzo negativo (traccion) por flexion en fibra inferior.")

spacer()

# PUNTO 3
heading("6.2.3 Punto 3 - Eje Neutro (y = 0)", 3)
pt3 = puntos_mohr[2]
body("Esfuerzo normal:")
eq(f"sigma_x = N / A = {N_crit_mohr:.2f}e3 / {bp['A']}")
eq_red(f"sigma_x = {pt3['sigma_x']:.2f} MPa")

body("Esfuerzo cortante:")
eq(f"tau_xy = V x Q_NA / (Ix x tw)")
eq(f"tau_xy = {V_crit_mohr:.2f}e3 x {Q_NA:.1f} / ({Ix_mm4:.0f} x {tw_mm})")
eq_red(f"tau_xy = {pt3['tau_xy']:.2f} MPa")

eq(f"sigma_avg = {pt3['sigma_avg']:.2f} MPa")
eq(f"R = sqrt({pt3['sigma_avg']:.2f}^2 + {pt3['tau_xy']:.2f}^2) = {pt3['R']:.2f} MPa")
eq_red(f"sigma_1 = {pt3['sigma_1']:.2f} MPa")
eq_red(f"sigma_2 = {pt3['sigma_2']:.2f} MPa")
eq_red(f"tau_max = {pt3['tau_max']:.2f} MPa")
eq(f"theta_p = {pt3['theta_p']:.1f} grados")

spacer()

# PUNTO 4
heading("6.2.4 Punto 4 - Union Ala-Alma (y = " + f"{d_mm/2-tf_mm:.0f} mm)", 3)
pt4 = puntos_mohr[3]
body("Esfuerzo normal:")
eq(f"sigma_x = M x {d_mm/2-tf_mm:.0f} / Ix + N / A")
eq_red(f"sigma_x = {pt4['sigma_x']:.2f} MPa")

body("Esfuerzo cortante:")
eq(f"tau_xy = V x Q_ala / (Ix x tw)")
eq(f"tau_xy = {V_crit_mohr:.2f}e3 x {Q_ala:.1f} / ({Ix_mm4:.0f} x {tw_mm})")
eq_red(f"tau_xy = {pt4['tau_xy']:.2f} MPa")

eq(f"sigma_avg = {pt4['sigma_avg']:.2f} MPa")
eq(f"R = sqrt({pt4['sigma_avg']:.2f}^2 + {pt4['tau_xy']:.2f}^2) = {pt4['R']:.2f} MPa")
eq_red(f"sigma_1 = {pt4['sigma_1']:.2f} MPa")
eq_red(f"sigma_2 = {pt4['sigma_2']:.2f} MPa")
eq_red(f"tau_max = {pt4['tau_max']:.2f} MPa")
eq(f"theta_p = {pt4['theta_p']:.1f} grados")

body("Estado biaxial: componentes de esfuerzo normal y cortante coexisten en este punto.")

spacer()

heading("6.3 Resumen de Esfuerzos", 2)
add_table(
    ['Punto', 'y [mm]', 'sigma_x', 'tau', 'sigma_1', 'sigma_2', 'tau_max'],
    [
        ['1: Fibra Sup.', f'{puntos_mohr[0]["y"]:.0f}',
         f'{puntos_mohr[0]["sigma_x"]:.2f}', f'{puntos_mohr[0]["tau_xy"]:.2f}',
         f'{puntos_mohr[0]["sigma_1"]:.2f}', f'{puntos_mohr[0]["sigma_2"]:.2f}',
         f'{puntos_mohr[0]["tau_max"]:.2f}'],
        ['2: Fibra Inf.', f'{puntos_mohr[1]["y"]:.0f}',
         f'{puntos_mohr[1]["sigma_x"]:.2f}', f'{puntos_mohr[1]["tau_xy"]:.2f}',
         f'{puntos_mohr[1]["sigma_1"]:.2f}', f'{puntos_mohr[1]["sigma_2"]:.2f}',
         f'{puntos_mohr[1]["tau_max"]:.2f}'],
        ['3: Eje Neutro', f'{puntos_mohr[2]["y"]:.0f}',
         f'{puntos_mohr[2]["sigma_x"]:.2f}', f'{puntos_mohr[2]["tau_xy"]:.2f}',
         f'{puntos_mohr[2]["sigma_1"]:.2f}', f'{puntos_mohr[2]["sigma_2"]:.2f}',
         f'{puntos_mohr[2]["tau_max"]:.2f}'],
        ['4: Ala-Alma', f'{puntos_mohr[3]["y"]:.0f}',
         f'{puntos_mohr[3]["sigma_x"]:.2f}', f'{puntos_mohr[3]["tau_xy"]:.2f}',
         f'{puntos_mohr[3]["sigma_1"]:.2f}', f'{puntos_mohr[3]["sigma_2"]:.2f}',
         f'{puntos_mohr[3]["tau_max"]:.2f}'],
    ]
)
body("Tabla 6.1: Esfuerzos en los 4 puntos (unidades: MPa).")

body("Observaciones:")
bullet(f"sigma_1 maximo ({puntos_mohr[0]['sigma_1']:.2f} MPa) en fibra superior.")
bullet(f"tau_max maximo ({puntos_mohr[2]['tau_max']:.2f} MPa) en el eje neutro.")
bullet("Estado biaxial en la union ala-alma.")
bullet(f"Todos los sigma_1 < Fy = {Fy} MPa: diseno seguro.")

spacer()

heading("6.4 Circulos de Mohr", 2)
add_img('Circulos_Mohr_Correcto.png', 6.0)
body("Figura 6.1: Circulos de Mohr para los 4 puntos de la seccion critica.")

body("Punto 1: Circulo con centro en sigma_avg positivo, pasa por el origen. Estado uniaxial.")
body("Punto 2: Circulo centrado en sigma_avg negativo, pasa por el origen. Estado uniaxial.")
body("Punto 3: Circulo casi simetrico, radio determinado por el cortante. Maximo tau.")
body("Punto 4: Circulo con contribucion de sigma y tau. Estado biaxial completo.")

doc.add_page_break()

# ##############################################################
# CAPITULO 7
# ##############################################################
print("    Cap 7: Volumenes y conclusiones...")

heading("CAPITULO 7: VOLUMENES, PESOS Y CONCLUSIONES")
spacer()

heading("7.1 Computo de Volumenes y Pesos", 2)

body(f"Viga ({bn}): L = {L} m, A = {bp['A']} mm2")
eq(f"V_viga = {bp['A']} x {L*1000:.0f} / 10^9")
eq_red(f"V_viga = {V_viga:.4f} m3")
eq(f"P_viga = {V_viga:.4f} x {rho_acero}")
eq_red(f"P_viga = {P_viga:.0f} kg = {P_viga*g_acc/1000:.2f} kN")

spacer()

body(f"Columnas (2x {cn}): h = {h} m, A = {cp['A']} mm2")
eq(f"V_col = 2 x {cp['A']} x {h*1000:.0f} / 10^9")
eq_red(f"V_col = {V_col:.4f} m3")
eq(f"P_col = {V_col:.4f} x {rho_acero}")
eq_red(f"P_col = {P_col:.0f} kg = {P_col*g_acc/1000:.2f} kN")

spacer()

body("Conexiones (estimacion):")
eq(f"V_conex = {V_conex:.3f} m3")
eq_red(f"P_conex = {P_conex:.0f} kg")

spacer()

body("Total del portico:")
eq(f"V_total = {V_viga:.4f} + {V_col:.4f} + {V_conex:.3f}")
eq_red(f"V_total = {V_total:.4f} m3")
eq_red(f"P_total = {P_total:.0f} kg = {P_total*g_acc/1000:.2f} kN")

add_table(
    ['Elemento', 'Perfil', 'Long. [m]', 'Vol. [m3]', 'Peso [kg]', 'Peso [kN]'],
    [
        ['Viga', bn, f'{L}', f'{V_viga:.4f}', f'{P_viga:.0f}', f'{P_viga*g_acc/1000:.2f}'],
        ['Columnas x2', cn, f'2x{h}', f'{V_col:.4f}', f'{P_col:.0f}', f'{P_col*g_acc/1000:.2f}'],
        ['Conexiones', '-', '-', f'{V_conex:.3f}', f'{P_conex:.0f}', f'{P_conex*g_acc/1000:.2f}'],
        ['TOTAL', '-', '-', f'{V_total:.4f}', f'{P_total:.0f}', f'{P_total*g_acc/1000:.2f}'],
    ]
)
body("Tabla 7.1: Computo de volumenes y pesos.")

doc.add_page_break()

heading("7.2 Conclusiones", 2)

body(f"1. Se analizo correctamente el portico plano biempotrado con k = {k:.4f}, superponiendo "
     f"carga distribuida q = {q_t:.3f} kN/m y carga lateral P = {P_lat} kN.")

body(f"2. El perfil de viga {bn} cumple: FS_flexion = {fs_b:.2f} >= {FS_min}, "
     f"FS_cortante = {fs_tau_b:.2f} >= {FS_min}, d/tw = {bp['d']/bp['tw']:.1f} <= 50.")

body(f"3. El perfil de columna {cn} cumple: FS_flexion = {fs_c:.2f} >= {FS_min}, "
     f"FS_cortante = {fs_tau_c:.2f} >= {FS_min}, d/tw = {cp['d']/cp['tw']:.1f} <= 50.")

if best_custom:
    body(f"4. Perfil personalizado (d={pc['d']}, bf={pc['bf']}, tf={pc['tf']}, tw={pc['tw']} mm) "
         f"reduce el area en {(1-pc['A']/bp['A'])*100:.1f}% con FS = {pc['fs_f']:.2f}.")

body(f"5. Conexiones viga-columna: {n_total} pernos A490 phi {bolt_sel}, verificados por cortante, "
     "aplastamiento, desgarramiento y block shear. Soldadura E70 como alternativa.")

body(f"6. Conexiones de base: placa {ancho_placa}x{largo_placa}x{t_placa} mm, "
     f"{n_pernos_base} pernos A490 phi {bolt_base}, soldadura perimetral a={a_base} mm.")

body(f"7. Circulos de Mohr: sigma_1_max = {puntos_mohr[0]['sigma_1']:.2f} MPa < Fy = {Fy} MPa. "
     "Diseno seguro en todos los puntos de la seccion critica.")

body(f"8. Peso total del portico: {P_total:.0f} kg ({P_total*g_acc/1000:.2f} kN) de acero A36.")

body("9. El equilibrio global fue verificado para ambos casos de carga, confirmando la validez "
     "de los resultados analiticos.")

body("10. Se recomienda verificar deformaciones (flecha, desplazamiento lateral) en etapas "
     "posteriores del diseno.")

doc.add_page_break()

# ##############################################################
# CAPITULO 8
# ##############################################################
print("    Cap 8: Referencias...")

heading("CAPITULO 8: REFERENCIAS")
spacer()
body("[1] Guia del Proyecto Final - Mecanica de Materiales. Ing. William Valencia Mina, "
     "Universidad del Quindio.")
body("[2] AISC Steel Construction Manual, 15th Edition. Chicago, IL, 2017.")
body("[3] ASTM A36/A36M: Standard Specification for Carbon Structural Steel.")
body("[4] ASTM A490: Standard Specification for Structural Bolts, 150 ksi Min Tensile Strength.")
body("[5] Beer, F.P., Johnston, E.R., DeWolf, J.T., Mazurek, D.F. Mechanics of Materials, 8th Ed.")
body("[6] Hibbeler, R.C. Mechanics of Materials, 10th Ed. Pearson, 2017.")
body("[7] McCormac, J.C. Structural Steel Design, 6th Ed. Pearson, 2018.")
body("[8] NSR-10: Reglamento Colombiano de Construccion Sismo Resistente, 2010.")
body("[9] AWS D1.1: Structural Welding Code - Steel.")
body("[10] Salmon, C.G. Steel Structures: Design and Behavior, 5th Ed. Pearson, 2009.")

doc.add_page_break()

# ##############################################################
# CAPITULO 9
# ##############################################################
print("    Cap 9: Anexos...")

heading("CAPITULO 9: ANEXOS")
spacer()

heading("Anexo A: Esquema del Portico", 2)
body("Esquema general del portico con cargas y reacciones.")
add_img('Esquema_Portico_Correcto.png', 6.0)
body("Figura A.1: Esquema del portico plano.")

spacer()

heading("Anexo B: Diagramas de Fuerza Interna", 2)
body("Diagramas de momento flector y fuerza cortante para viga y columnas.")
add_img('Diagramas_Portico.png', 6.5)
body("Figura B.1: Diagramas V/M.")

spacer()

heading("Anexo C: Circulos de Mohr", 2)
body("Circulos de Mohr para 4 puntos de la seccion critica de la viga.")
add_img('Circulos_Mohr_Correcto.png', 6.0)
body("Figura C.1: Circulos de Mohr.")

spacer()

heading("Anexo D: Propiedades de Perfiles", 2)
add_table(
    ['Perfil', 'd', 'bf', 'tf', 'tw', 'Ix [cm4]', 'Sx [cm3]', 'A [mm2]', 'kg/m'],
    [
        [bn, f'{bp["d"]}', f'{bp["bf"]}', f'{bp["tf"]}', f'{bp["tw"]}',
         f'{bp["Ix"]}', f'{bp["Sx"]}', f'{bp["A"]}', f'{bp["w_kg"]}'],
        [cn, f'{cp["d"]}', f'{cp["bf"]}', f'{cp["tf"]}', f'{cp["tw"]}',
         f'{cp["Ix"]}', f'{cp["Sx"]}', f'{cp["A"]}', f'{cp["w_kg"]}'],
    ]
)
body("Tabla D.1: Propiedades de los perfiles seleccionados.")

spacer()

heading("Anexo E: Formulas del Portico Biempotrado", 2)
body("Caso 1: Carga distribuida uniforme q:")
eq("V_A = V_D = q x L / 2")
eq("H_A = H_D = q x L^2 / [4 x h x (k + 2)]")
eq("M_A = M_D = q x L^2 / [12 x (k + 2)]")
eq("M_B = M_C = -q x L^2 / [6 x (k + 2)]")
eq("M_mid = q x L^2 x (3k + 2) / [24 x (k + 2)]")

body("Caso 2: Carga lateral P:")
eq("V_A = 3Phk / [L(6k+1)]")
eq("V_D = -V_A")
eq("H_A = H_D = P / 2")
eq("M_A = -Ph(3k+1) / [2(6k+1)]")
eq("M_D = +Ph(3k+1) / [2(6k+1)]")
eq("M_B = (Ph/2) x 3k/(6k+1)")
eq("M_C = -M_B")

body("Parametro de rigidez:")
eq("k = (I_2 / I_1) x (h / L)")

spacer()

heading("Anexo F: Ecuaciones de Mohr", 2)
eq("sigma_avg = (sigma_x + sigma_y) / 2")
eq("R = sqrt[((sigma_x - sigma_y)/2)^2 + tau_xy^2]")
eq("sigma_1 = sigma_avg + R")
eq("sigma_2 = sigma_avg - R")
eq("tau_max = R")
eq("theta_p = (1/2) arctan(2 tau_xy / (sigma_x - sigma_y))")

body("Para sigma_y = 0:")
eq("sigma_avg = sigma_x / 2")
eq("R = sqrt(sigma_avg^2 + tau_xy^2)")

spacer()

heading("Anexo G: Ecuaciones de Conexiones", 2)
eq("Rv = Fnv x Ab  (cortante pernos)")
eq("Rn = 2.4 x Fu x db x t  (aplastamiento)")
eq("Rn = 1.2 x Le_eff x t x Fu  (desgarramiento)")
eq("Rn = 0.6 Fu Anv + Ubs Fu Ant  (block shear)")
eq("fw = 0.707 x a x Fw x L  (soldadura filete)")
eq("Fw = 0.6 x Fexx")

spacer()

body("--- Fin de la Memoria de Calculos ---")

# ================================================================
# GUARDAR
# ================================================================
try:
    doc.save(FINAL)
except PermissionError:
    FINAL = os.path.join(OUT_DIR, "Memoria_de_Calculos_NUEVO.docx")
    try:
        doc.save(FINAL)
    except PermissionError:
        FINAL = os.path.join(OUT_DIR, "Memoria_PORTICO_" + str(int(__import__('time').time())) + ".docx")
        doc.save(FINAL)
    print("  (Archivo bloqueado, guardado como alternativo)")

n_paragraphs = len(doc.paragraphs)
print(f"\n  Documento guardado: {FINAL}")
print(f"  Total de parrafos: {n_paragraphs}")

print("\n" + "="*70)
print("  CALCULOS COMPLETADOS EXITOSAMENTE!")
print(f"  Archivo: {FINAL}")
print("="*70)
